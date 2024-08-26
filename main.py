import customtkinter as ctk
from tkinter import messagebox,END
from tkinter import *
import sqlite3
import tkinter as tk
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.lib.utils import ImageReader
from datetime import datetime , date
import os
from docx import Document  # Ensure this import is included
from tkinter import filedialog
from threading import Thread

import time
import sys
import imageio

from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from tkinter import filedialog
from PIL import Image, ImageTk
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import webbrowser




# -------------------------- DATABASE
def create_or_alter_table():
    try:
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        cursor.execute('''CREATE TABLE IF NOT EXISTS patients
                          (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, age INTEGER, disease TEXT, phone TEXT, gender TEXT)''')

        # Add gender column to patients table if it does not exist
        cursor.execute("PRAGMA table_info(patients)")
        columns = [info[1] for info in cursor.fetchall()]
        if 'gender' not in columns:
            cursor.execute("ALTER TABLE patients ADD COLUMN gender TEXT")
        if 'status' not in columns:
            cursor.execute("ALTER TABLE patients ADD COLUMN status TEXT DEFAULT 'Pending'")   


        # Create the prescriptions table if it does not exist
        cursor.execute('''CREATE TABLE IF NOT EXISTS prescriptions
                          (id INTEGER PRIMARY KEY AUTOINCREMENT, patient_id INTEGER, prescription TEXT, date TEXT,
                          FOREIGN KEY(patient_id) REFERENCES patients(id))''')

        cursor.execute('''CREATE TABLE IF NOT EXISTS SaveMedicalHistoryChecked
                          (patient_id INTEGER, SavedMedicalHistory TEXT, date TEXT,
                          FOREIGN KEY(patient_id) REFERENCES patients(id))''')

        cursor.execute('''CREATE TABLE IF NOT EXISTS SaveDentalHistoryChecked
                          (patient_id INTEGER, SavedDentalHistory TEXT, date TEXT,
                          FOREIGN KEY(patient_id) REFERENCES patients(id))''')

        # Add date column to prescriptions table if it does not exist
        cursor.execute("PRAGMA table_info(prescriptions)")
        columns = [info[1] for info in cursor.fetchall()]
        if 'date' not in columns:
            cursor.execute("ALTER TABLE prescriptions ADD COLUMN date TEXT")

        # Create the medicines table if it does not exist
        cursor.execute('''CREATE TABLE IF NOT EXISTS medicines
                          (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT)''')

        # Create the payments table if it does not exist


        cursor.execute('''CREATE TABLE IF NOT EXISTS payments
                  (id INTEGER PRIMARY KEY AUTOINCREMENT, patient_id INTEGER, paid INTEGER, remaining INTEGER, total INTEGER, date TEXT,
                  FOREIGN KEY(patient_id) REFERENCES patients(id))''')

        # Create the general_dental_history table for common dental history information
        cursor.execute('''CREATE TABLE IF NOT EXISTS general_dental_history
                          (id INTEGER PRIMARY KEY AUTOINCREMENT, history TEXT)''')

        # Create the general_medical_history table for common medical history information
        cursor.execute('''CREATE TABLE IF NOT EXISTS general_medical_history
                          (id INTEGER PRIMARY KEY AUTOINCREMENT, history TEXT)''')


        cursor.execute('''CREATE TABLE IF NOT EXISTS treatment_planning
                          (id INTEGER PRIMARY KEY AUTOINCREMENT, history TEXT)''')

        # Create the treatment_planning table if it does not exist, linked to patients table
        
                # Create the custom_data table if it does not exist
        cursor.execute('''CREATE TABLE IF NOT EXISTS custom_data
                          (id INTEGER PRIMARY KEY AUTOINCREMENT, patient_id INTEGER, custom_text TEXT, date TEXT,
                          FOREIGN KEY(patient_id) REFERENCES patients(id))''')
        # Create or modify a table to store the checked state



        conn.commit()
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
    finally:
        conn.close()

def get_resource_path(relative_path):
    """ Get the absolute path to a resource, working for both dev and PyInstaller """
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

# Example usage in your generate_pdf function
logo_path = get_resource_path("Images/logo.png")
logo_path2 = get_resource_path("Images/FullNameLogo.png")
logo_path3 = get_resource_path("Images/doctors.png")
logo_path4 = get_resource_path("Images/adress.png")
logo_path5 = get_resource_path("Images/backside.jpeg")
logo_path6 = get_resource_path("Images/temp_logo.jpg")

# Global variables for histories
selected_medical_history = ""
selected_dental_history = ""
selected_treatment_history = ""

# Global list of selected medicines
selected_medicines = []


def draw_selected_items(c, x_label_start, y_start, font_size, line_spacing):
    current_y = y_start

    # Draw Medical History if available
    if selected_medical_history:
        c.setFont('Helvetica-Bold', font_size + 1)
        c.setFillColor(HexColor("#4e4c4c"))  # Dark Gray color
        lines = selected_medical_history.split('\n')
        for i, line in enumerate(lines):
            c.drawString(x_label_start - 40, current_y - 200 - (i * line_spacing), line)
        current_y -= 90 + len(lines) * line_spacing

    # Draw Dental History if available
    if selected_dental_history:
        c.setFont('Helvetica-Bold', font_size + 1)
        c.setFillColor(HexColor("#4e4c4c"))  # Dark Gray color
        lines2 = selected_dental_history.split('\n')
        for i, line2 in enumerate(lines2):
            c.drawString(x_label_start - 40, current_y - 350 - (i * line_spacing), line2)
        current_y -= 90 + len(lines2) * line_spacing

    # Draw Treatment Planning if available
    if selected_treatment_history:
        current_x = x_label_start + 350
        c.setFont('Helvetica-Bold', font_size + 1)
        c.setFillColor(HexColor("#4e4c4c"))  # Dark Gray color
        lines3 = selected_treatment_history.split('\n')
        for i, line3 in enumerate(lines3):
            c.drawString(current_x - 90, y_start - 550 - (i * line_spacing), line3)

    if selected_medicines:
        c.setFont('Helvetica-Bold', font_size + 2)
        c.setFillColor(HexColor("#4e4c4c"))  # Dark Gray color

        # Load the transparent image using PIL
        checkmark_image = Image.open("Images/urdutext.png")
        checkmark_image_reader = ImageReader(checkmark_image)
        image_width, image_height = 50, 20  # Set the desired size of the checkmark image

        for i, (medicine, var) in enumerate(selected_medicines.items()):
            line_y_position = y_start - 190 - (i * line_spacing)
            c.drawString(x_label_start + 110, line_y_position, medicine)

            # Append the checkmark image if the medicine is checked
            if var.get():
                c.drawImage(checkmark_image_reader, x_label_start + 250 + c.stringWidth(medicine) + 5,
                            line_y_position - 3, width=image_width, height=image_height, anchor="e")

def generate_pdf(patient_id):
    patient = fetch_patient(patient_id)
    if not patient:
        messagebox.showerror("Error", "Patient not found")
        return

    # Determine the base path where the script is located
    if getattr(sys, 'frozen', False):
        # Running in a PyInstaller bundle
        base_path = sys._MEIPASS
    else:
        # Running in a normal Python environment
        base_path = os.path.dirname(os.path.abspath(__file__))

    # Define the output directory as the directory of the script
    output_dir = base_path
    os.makedirs(output_dir, exist_ok=True)  # Ensure the directory exists

    # Generate a unique file name based on patient name and timestamp
    file_name = f"prescription_{patient[1]}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"

    # Define the full path for the PDF file
    pdf_path = os.path.join(output_dir, file_name)

    # Create a canvas object and save the PDF
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4

    # Define the path to the logo images
    logo_path = os.path.join(output_dir, "Images/logo.png")
    logo_path2 = os.path.join(output_dir, "Images/FullNameLogo.png")
    logo_path3 = os.path.join(output_dir, "Images/doctors.png")
    logo_path4 = os.path.join(output_dir, "Images/adress.png")

    # Draw the logo images on the first page
    c.drawImage(logo_path, 50, height - 95, width=100, preserveAspectRatio=True, mask='auto')
    c.drawImage(logo_path2, 140, height - 170, width=420, preserveAspectRatio=True, mask='auto')

    # Draw a horizontal line under the logos
    line_start_x = 50
    line_end_x = 560
    line_y = height - 105
    c.setStrokeColor(HexColor("#da2020"))
    c.setLineWidth(3)
    c.line(line_start_x, line_y, line_end_x, line_y)

    # Draw the doctor image
    c.drawImage(logo_path3, 70, height - 310, width=420, preserveAspectRatio=True, mask='auto')

    # Draw the current date and patient details
    font_size = 12
    line_spacing = 20
    x_label_start = 50
    y_start = height - 210
    current_date = date.today().strftime("%B %d, %Y")

    c.setFont('Helvetica-Bold', font_size)
    c.setFillColor(HexColor("#000000"))
    c.drawString(x_label_start, y_start - line_spacing, "Date: ")
    c.setFont('Helvetica', font_size)
    c.setFillColor(HexColor("#4e4c4c"))
    c.drawString(x_label_start + 35, y_start - line_spacing, current_date)

    c.setFont('Helvetica-Bold', font_size)
    c.setFillColor(HexColor("#000000"))
    c.drawString(x_label_start + 140, y_start - line_spacing, "Pt. Name: ")
    c.setFont('Helvetica', font_size)
    c.setFillColor("#4e4c4c")
    c.drawString(x_label_start + 200, y_start - line_spacing, patient[1])

    c.setFont('Helvetica-Bold', font_size)
    c.setFillColor(HexColor("#000000"))
    c.drawString(x_label_start + 360, y_start - line_spacing, "Age: ")
    c.setFont('Helvetica', font_size)
    c.setFillColor("#4e4c4c")
    c.drawString(x_label_start + 390, y_start - line_spacing, str(patient[2]))

    c.setFont('Helvetica-Bold', font_size)
    c.setFillColor(HexColor("#000000"))
    c.drawString(x_label_start + 440, y_start - line_spacing, "Ref: ")
    c.setFont('Helvetica', font_size)
    c.setFillColor("#4e4c4c")
    c.drawString(x_label_start + 480, y_start - line_spacing, f"{patient[0]:03}")

    # Add sections and lines
    c.setFont('Helvetica-Bold', font_size + 2)
    c.setFillColor(HexColor("#000000"))
    c.drawString(x_label_start, y_start - 40 - line_spacing, "Medical History")
    c.drawString(x_label_start, y_start - 300 - line_spacing, "Dental History")
    c.drawString(x_label_start + 300, y_start - 370 - line_spacing, "Treatment Planning")

    draw_selected_items(c, 100, 750, 12, 15)

    # Draw a vertical line next to the patient's name
    line_start_x = 180
    line_start_y = y_start - 24 * line_spacing - inch / 2
    line_end_y = y_start - line_spacing - inch / 2 + 10

    c.setLineWidth(1)
    c.setStrokeColor("black")
    c.line(line_start_x, line_start_y, line_start_x, line_end_y)

    c.drawImage(logo_path4, 0, height - 910, width=600, preserveAspectRatio=True, mask='auto')

    # Save the canvas
    c.showPage()
    c.save()

    # Wait for a short duration to ensure the file is completely saved
    time.sleep(1)

    # Ensure correct path formatting for webbrowser
    pdf_url = f'file:///{os.path.abspath(pdf_path)}'.replace('\\', '/')
    webbrowser.open_new(pdf_url)

    # Wait for a short duration before deletion
    time.sleep(10)
    try:
        os.remove(pdf_path)
        print(f'{pdf_path} has been deleted.')
    except OSError as e:
        print(f'Error: {e.strerror}')

    # Optionally, print the PDF (ensure the printer is connected)
    if os.system("sc query spooler") != 0:
        messagebox.showerror("Printer Error", "Printer is not connected or not available")
    else:
        os.startfile(pdf_path, "print")






def save_prescription(patient_id):
    prescription = '\n'.join(selected_medicines)
    date = datetime.now().strftime("%d-%m-%Y %I:%M %p")
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("INSERT INTO prescriptions (patient_id, prescription, date) VALUES (?, ?, ?)", (patient_id, prescription, date))

    conn.commit()
    conn.close()
    messagebox.showinfo("Saved", "Prescription saved successfully")
    show_patients()

def save_medical_checked_history(patient_id):
        SavedmedicalHistory = selected_medical_history
        date = datetime.now().strftime("%Y-%m-%d")
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("INSERT INTO SaveMedicalHistoryChecked (patient_id, SavedMedicalHistory, date) VALUES (?, ?, ?)", (patient_id, SavedmedicalHistory, date))
        conn.commit()
        conn.close()



def save_dental_checked_history(patient_id):
        SaveddentalHistory = selected_dental_history
        date = datetime.now().strftime("%Y-%m-%d")
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("INSERT INTO SaveDentalHistoryChecked (patient_id, SavedDentalHistory, date) VALUES (?, ?, ?)",
                       (patient_id, SaveddentalHistory, date))
        conn.commit()
        conn.close()

def fetch_medical_checked_history(patient_id):
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("SELECT SavedmedicalHistory FROM SaveMedicalHistoryChecked WHERE patient_id=?", (patient_id,))
    rows = cursor.fetchone()
    conn.close()
    return rows[0] if rows else "No medical history found."

def fetch_dental_checked_history(patient_id):
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("SELECT SavedDentalHistory FROM SaveDentalHistoryChecked WHERE patient_id=? ORDER BY date DESC", (patient_id,))
        saved_history = cursor.fetchone()
        conn.close()

        if saved_history:
            return saved_history[0].split('\n')
        return []

def save_prescription_and_generate_pdf(patient_id):
    # Save the prescription
    save_prescription(patient_id)
    save_medical_checked_history(patient_id)
    save_dental_checked_history(patient_id)
    # Generate the PDF
    generate_pdf(patient_id)


# def save_status(patient_id, status):
#     cursor.execute("UPDATE patients SET status = ? WHERE id = ?", (status, patient_id))
#     conn.commit()
#     messagebox.showinfo("Success", "Status updated successfully")

def generatebackpdf():
    try:
        # Determine the base path for saving the file
        if getattr(sys, 'frozen', False):
            # Running in a PyInstaller bundle
            base_path = sys._MEIPASS
        else:
            # Running in a normal Python environment
            base_path = os.path.dirname(os.path.abspath(__file__))

        # Define the output directory as the directory of the script
        output_dir = base_path
        os.makedirs(output_dir, exist_ok=True)  # Ensure the directory exists

        # Generate a unique file name based on patient name and timestamp
        patient_name = "default_patient"  # Replace with actual patient name variable
        file_name = f"prescription_{patient_name}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"

        # Define the full path for the PDF file
        pdf_path = os.path.join(output_dir, file_name)

        # Create a canvas object and save the PDF
        c = canvas.Canvas(pdf_path, pagesize=A4)
        width, height = A4

        # Draw logo image
        logo_image_path = os.path.join(base_path, "Images/backside.jpeg")
        if not os.path.exists(logo_image_path):
            print(f"Error: Logo image not found at {logo_image_path}")
            return
        logo_image = Image.open(logo_image_path).convert("RGB")
        logo_image = logo_image.resize((550, 600), Image.LANCZOS)
        temp_logo_path = os.path.join(base_path, "Images/temp_logo.jpg")
        logo_image.save(temp_logo_path)

        # Centering the logo image
        image_width, image_height = logo_image.size
        image_x = (width - image_width) / 2  # Center the image horizontally
        image_y = height - image_height  # Position image vertically 50 units from the top
        c.drawImage(temp_logo_path, image_x, image_y, width=image_width, height=image_height)



        # Draw text information
        treatment = treatment_entry.get()
        date = date_entry.get()
        time = time_entry.get()

        # Centering the text
        c.setFont("Helvetica-Bold", 18)
        c.setFillColor(HexColor("#db3333"))

        # Title text
        title_text = "Next Appointment"
        title_width = c.stringWidth(title_text, "Helvetica-Bold", 22)
        c.drawString((width - title_width) / 2, image_y - 20, title_text)

        # Treatment text
        c.setFont("Helvetica", 14)
        treatment_text = f"Treatment:"
        treatment_width = c.stringWidth(treatment_text, "Helvetica-Bold", 18)
        c.drawString(50,  image_y - 60, treatment_text)
        c.setFont("Helvetica", 14)
        c.setFillColor(("black"))
        treatment_text = f"{treatment}"
        treatment_width = c.stringWidth(treatment_text, "Helvetica", 18)
        c.drawString(50,  image_y - 80, treatment_text)

        # Date text
        c.setFillColor(HexColor("#db3333"))
        date_text = f"Date:"
        date_width = c.stringWidth(date_text, "Helvetica-Bold", 18)
        c.drawString(350 ,  image_y - 60, date_text)
        c.setFillColor(("black"))
        date_text = f"{date}"
        date_width = c.stringWidth(date_text, "Helvetica", 18)
        c.drawString(350 ,  image_y - 80, date_text)

        # Time text
        c.setFillColor(HexColor("#db3333"))        
        time_text = f"Time:"
        time_width = c.stringWidth(time_text, "Helvetica-Bold", 18)
        c.drawString(350, image_y - 150, time_text)
        c.setFillColor(("black"))
        time_text = f"{time:}"
        time_width = c.stringWidth(time_text, "Helvetica", 18)
        c.drawString(350, image_y - 170, time_text)
        

        c.save()

        print(f"PDF successfully created at {pdf_path}")

        # Open PDF in the browser
        webbrowser.open(f'file://{pdf_path}', new=2)

        # Delete PDF after 5 seconds
        root.after(5000, lambda: delete_pdf(pdf_path))

    except Exception as e:
        print(f"An error occurred: {e}")

def delete_pdf(pdf_path):
    try:
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
            print(f"PDF successfully deleted at {pdf_path}")
        else:
            print(f"PDF already deleted or not found at {pdf_path}")
    except Exception as e:
        print(f"Error deleting PDF: {e}")    
# ------------------------------------------ MAIN WINDOW -------------------------------------
def create_dental_care_window():

    for widget in root.winfo_children():
        widget.destroy()
    # Set up the main frame to cover the entire root window
    main_frame = ctk.CTkFrame(root , fg_color="white")
    main_frame.pack(expand=True, fill="both")

    main_frame.columnconfigure((2, 3), weight=2)
    main_frame.columnconfigure(0, weight=1)
    main_frame.columnconfigure(1, weight=3)
    main_frame.rowconfigure((0, 1, 2 , 3, 4), weight=1)
    root.attributes('-fullscreen' , False)


    # Load and display logo at the top center
    logo_image = Image.open("Images/FullLogo.png")
    logo_image = logo_image.resize((880, 150), Image.LANCZOS)  # Resize the image (width, height)

    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = ctk.CTkLabel(main_frame, text="", image=logo_photo)
    logo_label.grid(row=0, column=0, columnspan=4, pady=10, sticky="n")

    line_frame = ctk.CTkFrame(main_frame, height=3, fg_color="#db3333")  # Set line height and color
    line_frame.grid(row=1, column=0,padx = 80 , columnspan=4, sticky="new")  # Add space below the line # Adjust the y-coordinate and color as needed

    # Contact information
    contact_label = ctk.CTkLabel(main_frame ,  text_color="white", font=("dubai" , 18 , "bold") , text="5 Umer Road Islampura Bazar Lahore\nWhatsapp: 92-335-9999375", fg_color="#db3333")
    contact_label.grid(row=4, column=1,  columnspan=4,padx = (0 , 10) , pady= 5, sticky="wes")
    #  Adjust the y-coordinate and color as needed

    # Load images
    # Load images
    images = [Image.open(f"Images/img{i}.png").resize((450, 800)) for i in range(1, 5)]
    photo_images = [ImageTk.PhotoImage(img) for img in images]
    image_label = ctk.CTkLabel(main_frame, image=photo_images[0], text=" ")
    image_label.grid(row=1, column=0, rowspan=4, padx=(10,10), pady=5, sticky="wnes")
    
    # Set up separate labels and buttons

    button1 = ctk.CTkButton(main_frame, text_color="white", command=show_patients,   font=("dubai" , 28 , "bold") ,text="Patient List")
    button1.grid(row=1, column=1, padx=10, pady=(60, 10), sticky="news")

    button2 = ctk.CTkButton(main_frame,text_color="white",  font=("dubai" , 28 , "bold") , command=add_details ,  text="Add Details")
    button2.grid(row=1, column=2, padx=10, pady=(60, 10),  sticky="news")


    button3 = ctk.CTkButton(main_frame, text_color="white", font=("dubai" , 28 , "bold") ,command=show_add_patient_form ,   text="Add New Patients")
    button3.grid(row=2, column=1, columnspan= 2,  padx=10, pady=10, sticky="news")

    button4 = ctk.CTkButton(main_frame, text_color="white", font=("dubai" , 28 , "bold") , text="Welcome Window", command=initialize_welcome_window )
    button4.grid(row=3, column=1, padx=10, pady=10, sticky="news")

    # Load and display logo at the top center
    logo_image2 = Image.open("Images/teeth.png")
    # logo_image2 = logo_image2.resize((), Image.LANCZOS)  # Resize the image (width, height)
    logo_photo2 = ImageTk.PhotoImage(logo_image2)
    logo_label2 = ctk.CTkLabel(main_frame, text="", image=logo_photo2)
    logo_label2.grid(row=3, column=2, columnspan=1, pady=10, sticky="news")

    

    def slide_image(current_img, next_img, step=0):
     if step <= 20:
        # Blend current and next image
        alpha = step / 20
        blended = Image.blend(current_img, next_img, alpha)
        blended_image = ImageTk.PhotoImage(blended)
        image_label.configure(image=blended_image)
        image_label.image = blended_image
        root.after(50, slide_image, current_img, next_img, step + 1)
     else:
        image_label.configure(image=next_img)

    def update_image(image_index=0):
     current_img = images[image_index]
     next_index = (image_index + 1) % len(images)
     next_img = images[next_index]
    
     next_photo_img = ImageTk.PhotoImage(next_img)
     photo_images[next_index] = next_photo_img
    
     slide_image(current_img, next_img)
     root.after(3000, update_image, next_index)

    update_image()
  

def fetch_patients():
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name, age, disease, phone, gender , status FROM patients")
    rows = cursor.fetchall()
    conn.close()
    return rows

# Function to save custom data



def fetch_medicines():
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM medicines")
    rows = cursor.fetchall()
    conn.close()
    return rows

# Function to fetch custom data for a patient



# Function to add medical history to the database
def add_medical_history(new_history):
    try:
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        # Ensure the table exists
        cursor.execute("CREATE TABLE IF NOT EXISTS general_medical_history (id INTEGER PRIMARY KEY AUTOINCREMENT, history TEXT)")

        # Fetch existing history
        cursor.execute("SELECT history FROM general_medical_history WHERE id = 1")
        row = cursor.fetchone()

        if row:
            # Update existing record
            updated_history = f"{row[0]}, {new_history}"
            cursor.execute("UPDATE general_medical_history SET history = ? WHERE id = 1", (updated_history,))
        else:
            # Insert new record
            cursor.execute("INSERT INTO general_medical_history (history) VALUES (?)", (new_history,))

        conn.commit()
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
    finally:
        conn.close()

# Function to fetch and display medical history
def fetch_medical_history():
    try:
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        cursor.execute("SELECT history FROM general_medical_history WHERE id = 1")
        row = cursor.fetchone()

        if row:
            return row[0]
        else:
            return "No medical history found."
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
        return "Error fetching data."
    finally:
        conn.close()

# Function to update the display frame with current data
def update_display_frame():
    all_history = fetch_medical_history()

    # Clear the frame before adding new content
    for widget in display_frame.winfo_children():
        widget.destroy()

    # Split the history into individual entries
    history_list = all_history.split(", ")
    headLabel = ctk.CTkLabel(display_frame , text_color="#db3333" , font=("dubai" , 18 , "bold") , text="Medical History List")
    headLabel.pack(pady=10)
    # Display each entry in a separate label, arranged vertically
    for entry in history_list:
        label_display = ctk.CTkLabel(display_frame,text_color="#7b847d" , font=("dubai" , 16 , "bold") , text=entry)
        label_display.pack(padx=10, pady=5, anchor="w") 

# Function to handle saving medical history from user input
def save_medical_history():
    new_history = entry_history.get()
    if new_history.strip():  # Check if input is not empty
        add_medical_history(new_history)
        update_display_frame()
        label_result.config(text="Saved successfully!", fg_color="green")
        entry_history.delete(0, 'end')  # Clear the entry field
    else:
        label_result.config(text="Please enter some text.", fg_color="red")



# dental history 
def add_dental_history(new_history):
    try:
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        # Ensure the table exists
        cursor.execute("CREATE TABLE IF NOT EXISTS general_dental_history (id INTEGER PRIMARY KEY AUTOINCREMENT, history TEXT)")

        # Fetch existing history
        cursor.execute("SELECT history FROM general_dental_history WHERE id = 1")
        row = cursor.fetchone()

        if row:
            # Update existing record
            updated_history = f"{row[0]}, {new_history}"
            cursor.execute("UPDATE general_dental_history SET history = ? WHERE id = 1", (updated_history,))
        else:
            # Insert new record
            cursor.execute("INSERT INTO general_dental_history (history) VALUES (?)", (new_history,))

        conn.commit()
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
    finally:
        conn.close()

# Function to fetch and display medical history
def fetch_dental_history():
    try:
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        cursor.execute("SELECT history FROM general_dental_history WHERE id = 1")
        row = cursor.fetchone()

        if row:
            return row[0]
        else:
            return "No dental history found."
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
        return "Error fetching data."
    finally:
        conn.close()

# Function to update the display frame with current data
def update_display_frame2():
    all_history2 = fetch_dental_history()

    # Clear the frame before adding new content
    for widget in display_frame2.winfo_children():
        widget.destroy()

    # Split the history into individual entries
    history_list2 = all_history2.split(", ")
    headLabel2 = ctk.CTkLabel(display_frame2 , text_color="#db3333" , font=("dubai" , 18 , "bold") , text="Dental History List")
    headLabel2.pack(pady=10)
    # Display each entry in a separate label, arranged vertically
    for entry in history_list2:
        label_display = ctk.CTkLabel(display_frame2,text_color="#7b847d" , font=("dubai" , 16 , "bold") , text=entry)
        label_display.pack(padx=10, pady=5, anchor="w") 



# Function to handle saving medical history from user input
def save_dental_history():
    new_history2 = entry_history2.get()
    if new_history2.strip():  # Check if input is not empty
        add_dental_history(new_history2)
        update_display_frame2()
        label_result.config(text="Saved successfully!", fg_color="green")
        entry_history2.delete(0, 'end')  # Clear the entry field
    else:
        label_result.config(text="Please enter some text.", fg_color="red")





# Treatment Planning history 
def add_treatment_planning(new_history):
    try:
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        # Ensure the table exists
        cursor.execute("CREATE TABLE IF NOT EXISTS treatment_planning (id INTEGER PRIMARY KEY AUTOINCREMENT, history TEXT)")

        # Fetch existing history
        cursor.execute("SELECT history FROM treatment_planning WHERE id = 1")
        row = cursor.fetchone()

        if row:
            # Update existing record
            updated_history = f"{row[0]}, {new_history}"
            cursor.execute("UPDATE treatment_planning SET history = ? WHERE id = 1", (updated_history,))
        else:
            # Insert new record
            cursor.execute("INSERT INTO treatment_planning (history) VALUES (?)", (new_history,))

        conn.commit()
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
    finally:
        conn.close()

# Function to fetch and display medical history
def fetch_treatment_planning():
    try:
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        cursor.execute("SELECT history FROM treatment_planning WHERE id = 1")
        row = cursor.fetchone()

        if row:
            return row[0]
        else:
            return "No treatment planning history found."
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
        return "Error fetching data."
    finally:
        conn.close()

# Function to update the display frame with current data
def update_display_frame3():
    all_history = fetch_treatment_planning()

    # Clear the frame before adding new content
    for widget in display_frame3.winfo_children():
        widget.destroy()

    # Split the history into individual entries
    history_list = all_history.split(", ")
    headLabel = ctk.CTkLabel(display_frame3 , text_color="#db3333" , font=("dubai" , 18 , "bold") , text="Treatment Planning List")
    headLabel.pack(pady=10)
    # Display each entry in a separate label, arranged vertically
    for entry in history_list:
        label_display = ctk.CTkLabel(display_frame3,text_color="#7b847d" , font=("dubai" , 16 , "bold") , text=entry)
        label_display.pack(padx=10, pady=5, anchor="w") 

# Function to handle saving medical history from user input
def save_treatment_planning():
    new_history3 = entry_history3.get()
    if new_history3.strip():  # Check if input is not empty
        add_treatment_planning(new_history3)
        update_display_frame3()
        label_result.config(text="Saved successfully!", fg_color="green")
        entry_history.delete(0, 'end')  # Clear the entry field
    else:
        label_result.config(text="Please enter some text.", fg_color="red")

# Set up the GUI layout
def add_details():
    global entry_history, label_result, display_frame, display_frame2, entry_history2, display_frame3, entry_history3, notebook

    for widget in root.winfo_children():
        widget.destroy()

    main_frame = Frame(root)
    main_frame.pack(fill=BOTH, expand=True)
    main_frame.columnconfigure((0, 1), weight=1)
    main_frame.rowconfigure((0, 1, 2), weight=1)

    logo_image = Image.open("Images/FullLogo.png")
    logo_image = logo_image.resize((880, 160), Image.LANCZOS)
    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = ctk.CTkLabel(main_frame, text="", image=logo_photo)
    logo_label.grid(row=0, column=0, columnspan=2, pady=10, sticky="n")

    line_frame = ctk.CTkFrame(main_frame, height=3, fg_color="#db3333")
    line_frame.grid(row=1, column=0, columnspan=2, padx=80, sticky="new")

    notebook = ctk.CTkTabview(main_frame, segmented_button_selected_color="#db3333" , segmented_button_fg_color="#b0f3ff" ,
                            fg_color="#b0f3ff", width=800, height=400, anchor="w", corner_radius=5 , 
                            segmented_button_unselected_color="#05fa9a" )
    notebook.grid(row=1, column=0, columnspan=2 , pady = 20)
    notebook.columnconfigure((0, 1), weight=1)
    notebook.rowconfigure((0, 1, 2), weight=1)

    # TAB 1
    tab1 = notebook.add("Enter Medical History")
    tab1.configure(fg_color="#b0f3ff")  # Set tab 1 color to pink

    label_history = ctk.CTkLabel(tab1,text_color="#db3333" , font=("dubai" , 22 , "bold") , text="Enter Medical History")
    label_history.grid(column=1, row=1, sticky="ewn", columnspan=1, padx=150, pady=50)

    entry_history = ctk.CTkEntry(tab1, width=200, height=50)
    entry_history.grid(column=1, row=1, sticky="new", columnspan=1, padx=150, pady=100)

    button_save = ctk.CTkButton(tab1,fg_color="#00d6ff" , font=("dubai" , 16 , "bold") , text="Save Medical History", command=save_medical_history)
    button_save.grid(column=1, row=1, columnspan=1, padx=150, pady=150)

    entry_history.bind("<Return>", lambda event: save_medical_history())

    display_frame = ctk.CTkScrollableFrame(tab1, height=400, fg_color="white")
    display_frame.grid(column=0, row=0, rowspan=4, sticky="w")

    update_display_frame()

    # TAB 2
    tab2 = notebook.add("Enter Dental History")
    tab2.configure(fg_color="#b0f3ff")  # Set tab 2 color to green

    label_history2 = ctk.CTkLabel(tab2,text_color="#db3333" , font=("dubai" , 22 , "bold") , text="Enter Dental History")
    label_history2.grid(column=1, row=1, sticky="ewn", columnspan=1, padx=150, pady=50)

    entry_history2 = ctk.CTkEntry(tab2, width=200, height=50)
    entry_history2.grid(column=1, row=1, sticky="new", columnspan=1, padx=150, pady=100)

    button_save2 = ctk.CTkButton(tab2,fg_color="#00d6ff" , font=("dubai" , 16 , "bold") , text="Save Dental History", command=save_dental_history)
    button_save2.grid(column=1, row=1, columnspan=1, padx=150, pady=150)

    entry_history2.bind("<Return>", lambda event: save_dental_history())

    display_frame2 = ctk.CTkScrollableFrame(tab2, height=400, fg_color="white")
    display_frame2.grid(column=0, row=0, rowspan=4, sticky="w")

    update_display_frame2()

    # TAB 3
    tab3 = notebook.add("Enter Treatment Planning")
    tab3.configure(fg_color="#b0f3ff")  # Set tab 2 color to green

    label_history3 = ctk.CTkLabel(tab3,text_color="#db3333" , font=("dubai" , 22 , "bold") , text="Enter Treatment Planning ")
    label_history3.grid(column=1, row=1, sticky="ewn", columnspan=1, padx=150, pady=50)

    entry_history3 = ctk.CTkEntry(tab3, width=200, height=50)
    entry_history3.grid(column=1, row=1, sticky="new", columnspan=1, padx=150, pady=100)

    button_save3 = ctk.CTkButton(tab3,fg_color="#00d6ff" , font=("dubai" , 16 , "bold") , text="Save Treatment Planning", command=save_treatment_planning)
    button_save3.grid(column=1, row=1, columnspan=1, padx=150, pady=150)
    entry_history3.bind("<Return>", lambda event: save_treatment_planning())
    display_frame3 = ctk.CTkScrollableFrame(tab3, height=400, fg_color="white")
    display_frame3.grid(column=0, row=0, rowspan=4, sticky="w")

    update_display_frame3()






    ctk.CTkButton(tab1, text="Back", font=("dubai", 13, "bold"), fg_color="#ff79c1", hover_color="red", width=150, command=create_dental_care_window).grid(row=3, column=1, columnspan=1, pady=10)
    ctk.CTkButton(tab2, text="Back", font=("dubai", 13, "bold"), fg_color="#ff79c1", hover_color="red", width=150, command=create_dental_care_window).grid(row=3, column=1, columnspan=1, pady=10)
    ctk.CTkButton(tab3, text="Back", font=("dubai", 13, "bold"), fg_color="#ff79c1", hover_color="red", width=150, command=create_dental_care_window).grid(row=3, column=1, columnspan=1, pady=10)




 




def confirm_delete_patient(patient_id):
    response = messagebox.askyesno("Confirm Deletion", "Are you sure you want to delete this patient?")
    if response:
        delete_patient(patient_id)

def show_add_patient_form():
    for widget in root.winfo_children():
        widget.destroy()
        
    behindframe = ctk.CTkFrame(root , fg_color="#fddef5")    
    behindframe.pack(expand = 1, fill= 'both')


    # Create a main frame with larger width and height
    main_frame = ctk.CTkFrame(behindframe, width=1600, height=900, fg_color="white", corner_radius=20, border_color="#db3333", border_width=2)
    main_frame.place(relx=0.5, rely=0.5, anchor="center")

    # Define column weights
    main_frame.columnconfigure(0, weight=1)
    main_frame.columnconfigure(1, weight=1)

    # Left frame for the login form
    left_frame = ctk.CTkFrame(main_frame, fg_color="white", corner_radius=20)
    left_frame.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")

# Load and display the first image
# Load, resize, and display the first image
    img1 = Image.open("Images/LogoWithNH.png")
    img1_resized = resize_image(img1, width=150, height=150)
    img1 = ImageTk.PhotoImage(img1_resized)
    labelImage1 = ctk.CTkLabel(left_frame, text="", image=img1)
    labelImage1.grid(row=0, column=0 , pady = 10)
    labelImage1.image = img1



    # # Add Login label
    # login_label = ctk.CTkLabel(left_frame, text_color="#db3333", text="Welcome", font=("Dubai", 32, "bold"))
    # login_label.grid(row=0, column=0, columnspan=2, pady=10,padx=(10, 0))

    # Add subtitle label
    subtitle_label = ctk.CTkLabel(left_frame,text_color="#db3333", text="Add New Patient", font=("Dubai", 24, "bold"))
    subtitle_label.grid(row=1, column=0, columnspan=2, pady=15 , padx=(10, 0))

    # Username entry
    username_label = ctk.CTkLabel(left_frame, text="Name", font=("Dubai", 14, "bold"))
    username_label.grid(row=2, column=0, padx=10, pady=10, sticky="w")
    global name_entry
    name_entry = ctk.CTkEntry(left_frame,font=("Dubai", 14, "bold") ,width=200, placeholder_text="Enter Name" , placeholder_text_color="gray")
    name_entry.grid(row=2, column=1, padx=10, pady=10)

    # Password entry
    password_label = ctk.CTkLabel(left_frame, text="Age", font=("Dubai", 14, "bold"))
    password_label.grid(row=3, column=0, padx=10, pady=10, sticky="w")
    global age_entry
    age_entry = ctk.CTkEntry(left_frame,font=("Dubai", 14, "bold") ,width=200, placeholder_text="Enter Age" , placeholder_text_color="gray")
    age_entry.grid(row=3, column=1, padx=10, pady=10)
    # Username entry
    username_label = ctk.CTkLabel(left_frame, text="Treatment Protocol", font=("Dubai", 14, "bold"))
    username_label.grid(row=4, column=0, padx=10, pady=10, sticky="w")
    global disease_entry
    disease_entry = ctk.CTkEntry(left_frame,font=("Dubai", 14, "bold") ,width=200, placeholder_text="Enter Treatment Protocol " , placeholder_text_color="gray")
    disease_entry.grid(row=4, column=1, padx=10, pady=10)

    # Password entry
    password_label = ctk.CTkLabel(left_frame, text="Phone Number", font=("Dubai", 14, "bold"))
    password_label.grid(row=5, column=0, padx=10, pady=10, sticky="w")
    global phone_entry
    phone_entry = ctk.CTkEntry(left_frame,font=("Dubai", 14, "bold") ,width=200, placeholder_text="Enter Phone Number" , placeholder_text_color="gray")
    phone_entry.grid(row=5, column=1, padx=10, pady=10)

    # Gender selection using a dropdown menu
    gender_label = ctk.CTkLabel(left_frame, text="Gender", font=("Dubai", 14, "bold"))
    gender_label.grid(row=6, column=0, padx=10, pady=10, sticky="w")
    global gender_var
    gender_var = ctk.StringVar(value="Male")  # Default value is "Male"
    gender_dropdown = ctk.CTkOptionMenu(left_frame, fg_color="#db3333", button_color="#db3333",  variable=gender_var, values=["Male", "Female", "Other"])
    gender_dropdown.grid(row=6, column=1, padx=10, pady=10)

    # Save button
    # save_button = ctk.CTkButton(left_frame, text="Save", font=(

# Binding the update function to changes in gender_var
    # gender_var.trace("w", lambda *args: update_gender_text())


    # Login button
    login_button = ctk.CTkButton(left_frame, text="Save", font=("dubai" , 13 , "bold") ,  command=save_patient_info, fg_color="#db3333", hover_color="red", width=150)
    login_button.grid(row=7, column=0, columnspan=2, pady=20)
    ctk.CTkButton(left_frame, text="Back to Patient List",font=("dubai" , 13 , "bold") ,fg_color="#ff79c1", hover_color="red", width=150,command=show_patients).grid(row=8, column=0, columnspan=2,  pady=10)


    # Right frame for the image
    right_frame = ctk.CTkFrame(main_frame, fg_color="#D1C4E9", corner_radius=20)
    right_frame.grid(row=0, column=1, padx=20, pady=20, sticky="nsew")

    # Load and resize the image to fit the right frame
    image = Image.open("Images/NPT.jpg")



    def update_image():
        frame_width = right_frame.winfo_width()+40
        frame_height = right_frame.winfo_height()

        if frame_width <= 0 or frame_height <= 0:
            return  # Return if dimensions are not yet available

        # Resize the image to match the frame's height while maintaining aspect ratio
        image_width, image_height = image.size
        new_height = frame_height
        new_width = int(image_width * (new_height / image_height))

        image_resized = image.resize((new_width, new_height), Image.LANCZOS)
        photo = ImageTk.PhotoImage(image_resized)

        # Place the image in the right frame
        image_label = ctk.CTkLabel(right_frame, image=photo, text="")
        image_label.place(relx=0.5, rely=0.5, anchor="center")

        # Keep a reference to the image to prevent it from being garbage collected
        image_label.image = photo

    # Bind the update function to frame size changes
    right_frame.bind("<Configure>", lambda event: update_image())

def resize_image(image, width=None, height=None):
    original_width, original_height = image.size
    if width is not None and height is not None:
        ratio = min(width / original_width, height / original_height)
        new_width = int(original_width * ratio)
        new_height = int(original_height * ratio)
    elif width is not None:
        ratio = width / original_width
        new_width = width
        new_height = int(original_height * ratio)
    elif height is not None:
        ratio = height / original_height
        new_width = int(original_width * ratio)
        new_height = height
    else:
        new_width, new_height = original_width, original_height
    
    return image.resize((new_width, new_height), Image.LANCZOS)


def save_patient_info():
    name = name_entry.get()
    age = age_entry.get()
    disease = disease_entry.get()
    phone = phone_entry.get()
    gender = gender_var.get()

    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("INSERT INTO patients (name, age, disease, phone, gender) VALUES (?, ?, ?, ?, ?)", (name, age, disease, phone, gender))

    conn.commit()
    conn.close()

    messagebox.showinfo("Saved", "Patient information saved successfully")
    show_patients()

def delete_patient(patient_id):
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM patients WHERE id=?", (patient_id,))
    conn.commit()
    conn.close()
    show_patients()


def save_to_file():
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt"), ("All files", "*.*")])
    if file_path:
        with open(file_path, 'w') as file:
            for prescription in prescriptions:
                file.write(f"Prescription-ID: {prescription[0]:03}\n")
                file.write(f"Date: {prescription[3][:10]}\n")
                file.write(f"Time: {prescription[3][11:]}\n")
                file.write(f"Prescription: {prescription[2]}\n")
                file.write("\n\n\n")
    
    # Adding a right border to the left frame
def delete_patient(patient_id):
    confirmation = messagebox.askyesno("Delete Patient", "Are you sure you want to delete this patient permanently?")
    if confirmation:
        try:
            conn = sqlite3.connect('patients.db')
            cursor = conn.cursor()
            cursor.execute("DELETE FROM patients WHERE id=?", (patient_id,))
            cursor.execute("DELETE FROM prescriptions WHERE patient_id=?", (patient_id,))
            cursor.execute("DELETE FROM payments WHERE patient_id=?", (patient_id,))
            conn.commit()
        except sqlite3.Error as e:
            print(f"An error occurred: {e}")
        finally:
            conn.close()
        messagebox.showinfo("Deleted", "Patient has been deleted successfully.")
        show_patients()  # Refresh the patient list after deletion

def save_status(patient_id, status):
    try:
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        cursor.execute('''UPDATE patients SET status = ? WHERE id = ?''', (status, patient_id))
        conn.commit()
        
        messagebox.showinfo("Success", "Status updated successfully!")
    except sqlite3.Error as e:
        print(f"An error occurred: {e}")
        messagebox.showerror("Error", "An error occurred while updating status.")
    finally:
        conn.close()

def HistoryPage(patient_id):
    for widget in root.winfo_children():
        widget.destroy()

    main_frame = ctk.CTkFrame(root)
    main_frame.pack(fill=BOTH, expand=1)

    patient = fetch_patient(patient_id)
    global prescriptions
    prescriptions = fetch_prescriptions(patient_id)


    left_frame = ctk.CTkFrame(main_frame , fg_color="white")
    left_frame.pack(side="left", fill=Y)



    

    logo_image = Image.open("Images/logo.png")
    logo_image = logo_image.resize((175, 130), Image.LANCZOS)  # Resize the image (width, height)
    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = ctk.CTkLabel(left_frame, text="", image=logo_photo)
    logo_label.pack(padx=30, pady=40)

    right_frame = ctk.CTkScrollableFrame(main_frame , fg_color="white")
    right_frame.pack(fill=BOTH , expand=1)
    right_frame.columnconfigure((0, 1, 2, 3), weight=1)

    logo_image2 = Image.open("Images/FullNameLogo.png")
    logo_image2 = logo_image2.resize((800, 110), Image.LANCZOS)  # Resize the image (width, height)
    logo_photo2 = ImageTk.PhotoImage(logo_image2)
    logo_label2 = ctk.CTkLabel(right_frame, text="", image=logo_photo2)
    logo_label2.grid(row=0, column=0, columnspan=4, padx=30, pady=40)

    # Add a separator line
    line_frame2 = ctk.CTkFrame(right_frame, height=3, fg_color="#db3333")
    line_frame2.grid(row=1, column=0, columnspan=4, padx=30, sticky="new")

    notebook = ctk.CTkTabview(right_frame, fg_color="#c6f5fd", segmented_button_fg_color="#ffd9e4", 
                              segmented_button_selected_color="#db3333", segmented_button_selected_hover_color="red", 
                              segmented_button_unselected_color="#2dd293", segmented_button_unselected_hover_color="#0af579", 
                              width=800, height=500, anchor="w", corner_radius=25)
    notebook.grid(row=2, column=0, columnspan=4 , pady = 10)

    # TAB 1 - Prescription History
    tab1 = notebook.add("Prescription History")
    tab1.columnconfigure((0, 1, 2, 3), weight=1)

    # Display prescriptions
    for i, prescription in enumerate(prescriptions):
        ctk.CTkLabel(tab1, text=f"Prescription-ID: {prescription[0]:03}", font=("Dubai", 18, "bold"), 
                     text_color="#db3333").grid(row=0+i*3, column=0, pady=20)
        ctk.CTkLabel(tab1, text="Date:", font=("Dubai", 18, "bold"), text_color="#db3333").grid(row=0+i*3, column=2, 
                                                                                               sticky="e", pady=20)
        ctk.CTkLabel(tab1, text=f"{prescription[3][:10]}", font=("Dubai", 18), text_color="black").grid(row=0+i*3, 
                                                                                                         column=3, pady=20)
        ctk.CTkLabel(tab1, text="Time:", font=("Dubai", 18, "bold"), text_color="#db3333").grid(row=1+i*3, column=2, 
                                                                                               sticky="e", pady=20)
        ctk.CTkLabel(tab1, text=f"{prescription[3][11:]}", font=("Dubai", 18), text_color="black").grid(row=1+i*3, 
                                                                                                         column=3, pady=20)
        ctk.CTkLabel(tab1, text=f"{prescription[2]}", font=("Dubai", 18, "bold"), bg_color="#ffc9e4", 
                     height=300, width=500).grid(row=2+i*3, column=1, pady=20)
        
    medical_history_text = fetch_medical_checked_history(patient_id)
    dental_history_text = fetch_dental_checked_history(patient_id)

    # Assuming 'tab2' is already defined
    tab2 = notebook.add("Medical History")
    # tab2.columnconfigure((0, 1, 2, 3), weight=1)
    
    global history_label
    history_label = ctk.CTkLabel(tab2, text="", text_color="black", font=("Dubai", 18, "bold"))
    history_label.grid(row=0, column=0, columnspan=4, padx=10, pady=10, sticky="w")

        # Fetch saved medical history to display on initialization
        # patient_id = 1  # Example patient ID
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("SELECT SavedMedicalHistory, date FROM SaveMedicalHistoryChecked WHERE patient_id=? ORDER BY date DESC", (patient_id,))
    saved_history = cursor.fetchone()
    conn.close()

    if saved_history:
            selected_medical_history = saved_history[0]
            history_label.configure(text=f"{saved_history[1]}:\n{selected_medical_history}")
    else:
            selected_medical_history = ""
 
    tab3 = notebook.add("Dental History")
    tab3.columnconfigure((0, 1, 2, 3), weight=1)
    
    global history_label2
    history_label2 = ctk.CTkLabel(tab3, text="", text_color="black", font=("Dubai", 18, "bold"))
    history_label2.grid(row=0, column=0, columnspan=4, padx=10, pady=10, sticky="w")

        # Fetch saved dental history to display on initialization
        # patient_id = 1  # Example patient ID
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("SELECT SavedDentalHistory, date FROM SaveDentalHistoryChecked WHERE patient_id=? ORDER BY date DESC", (patient_id,))
    saved_history = cursor.fetchone()
    conn.close()

    if saved_history:
            selected_dental_history = saved_history[0]
            history_label2.configure(text=f"{saved_history[1]}:\n{selected_dental_history}")
    else:
            selected_dental_history = ""
 


    def save_custom_data(patient_id, custom_text):
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        current_date = datetime.now().strftime("%Y-%m-%d")
        cursor.execute("INSERT INTO custom_data (patient_id, custom_text, date) VALUES (?, ?, ?)", 
                   (patient_id, custom_text, current_date))
    
        conn.commit()
        conn.close()

    def load_custom_data(patient_id):
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()

        cursor.execute("SELECT custom_text, date FROM custom_data WHERE patient_id=?", (patient_id,))
        entries = cursor.fetchall()
        conn.close()

        return entries

    def manage_custom_data_tab(tab):
        def save_data():
            custom_text_value = custom_text.get("1.0", ctk.END).strip()
            save_custom_data(patient_id, custom_text_value)
            update_display()
    
        def update_display():
            entries = load_custom_data(patient_id)
            custom_data_display.configure(state="normal")
            custom_data_display.delete("1.0", ctk.END)

            for text, date in entries:
                custom_data_display.insert(ctk.END, f":\n{text}\n\n")
        
            custom_data_display.configure(state="disabled")

    # Create two frames inside tab4
        left_frame = ctk.CTkFrame(tab, width=400, height=400, corner_radius=10, fg_color="#f2f2f2")
        left_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)

        right_frame = ctk.CTkFrame(tab, width=400, height=400, corner_radius=10, fg_color="#f2f2f2")
        right_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)

    # Left frame for entry and save button
        ctk.CTkLabel(left_frame, text="Enter Custom Data:", font=("Dubai", 18, "bold")).pack(pady=10)
        custom_text = ctk.CTkTextbox(left_frame, width=300, height=200)
        custom_text.pack(pady=10)
    
        save_button = ctk.CTkButton(left_frame, text="Save Custom Data", command=save_data)
        save_button.pack(pady=10)

    # Right frame for displaying saved custom data
        ctk.CTkLabel(right_frame, text="Custom Data:", font=("Dubai", 18, "bold")).pack(pady=10)
        custom_data_display = ctk.CTkTextbox(right_frame, width=300, height=200, state="disabled")
        custom_data_display.pack(pady=10)

    # Load and display custom data
        update_display()


    tab4 = notebook.add("Custom Data")
    manage_custom_data_tab(tab4)





# report
    # Function to fetch patient details by ID or all patients if no ID is provided



        
    def fetch_patients(patient_id=None):
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        if patient_id:
            cursor.execute("SELECT id, name, age, disease, phone, gender , status FROM patients")

        else:
            cursor.execute("SELECT id, name, age, disease, phone, gender , status FROM patients")

        rows = cursor.fetchall()
        conn.close()
        return rows


    
    # Function to get specific patient details
    def get_patient_details(patient_id):
        patients = fetch_patients(patient_id)
        if not patients:
            return "No patient found with the provided ID."
        
        patient = patients[0]

        details = (
            f"ID:\t\t {patient[0]}\n"
            f"Name:\t\t {patient[1]}\n"
            f"Age:\t\t {patient[2]}\n"
            f"Disease:\t\t {patient[3]}\n"
            f"Phone:\t\t {patient[4]}\n"
            f"Gender:\t\t {patient[5]}\n"
        )
        return details

    # Function to get medical history
    def get_medical_history(patient_id):
        return fetch_medical_checked_history(patient_id)

    # Function to get dental history






    def get_dental_history(patient_id):
        dental_history_list = fetch_dental_checked_history(patient_id)
        return "\n".join(dental_history_list) if dental_history_list else "No dental history found."
    
    def fetch_custom_info(patient_id):
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("SELECT custom_text FROM custom_data WHERE patient_id=?", (patient_id,))
        rows = cursor.fetchall()
        conn.close()
        return "\n".join([row[0] for row in rows]) if rows else "No custom information found."

    # Function to get payment history (replace with actual implementation)
    def get_payment_history():
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("SELECT paid, remaining, total, date FROM payments WHERE patient_id=?", (patient_id,))
        payments = cursor.fetchall()
        conn.close()
    
        if not payments:
            return "No payment history available."

        content = ""
        for i, payment in enumerate(payments):
            content += (
            f"Payment Record-{i+1}:\n"
            f"Date:\t {payment[3]}\n"
            f"Paid:\t {payment[0]}\n"
            f"Remaining:\t {payment[1]}\n"
            f"Total:\t {payment[2]}\n\n"
        )
        return content
    
    def get_prescription_history(patient_id):
        PR = fetch_prescriptions(patient_id)
        if not PR:
            return "No prescription history available."

        content = ""
        for i, prescription in enumerate(PR):
            content += (
                f"Prescription-ID:\t\t {prescription[0]:03}\n"
                f"Date:\t\t {prescription[3][:10]}\n"
                f"Time:\t\t {prescription[3][11:]}\n"
                f"Prescription:\t\t {prescription[2]}\n\n"
            )
        return content

    # Function to generate the report
    def generate_report_for_patient():
        # Prompt for patient ID
       
        # Create the document
        doc = Document()
        doc.add_heading('Patient Information Report', level=1)
        
        # Retrieve data from tabs
        data = {
            'Patient Info': get_patient_details(patient_id),
            'Prescription History': get_prescription_history(patient_id),
            'Medical History': get_medical_history(patient_id),
            'Dental History': get_dental_history(patient_id),
            'Payment History': get_payment_history(),
            'Custom Info': fetch_custom_info(patient_id),


            # Add more sections if needed
        }
        
        for section, content in data.items():
            doc.add_heading(section, level=2)
            doc.add_paragraph(content)
        
        # Save the document using a file dialog
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if file_path:
            doc.save(file_path)
            print(f"Report generated successfully: {file_path}")

    tab5 = notebook.add('Report')  # Assuming you have other content in this tab
    # Add the Report button
    ctk.CTkLabel(tab5 ,font=("dubai" , 26 , "bold") , text_color="#db3333" ,  text="Generate  Report of you Patient").pack(pady = 60 , padx = 40)
    report_button = ctk.CTkButton(tab5, fg_color="#00c2ff" ,  text="Generate Report", command=generate_report_for_patient)
    report_button.pack(pady=10)





# Payments 

    def manage_payments():
        def calculate_remaining():
            total = int(total_entry.get() or 0)
            paid = int(paid_entry.get() or 0)
            remaining = total - paid
            remaining_label.configure(text=f"Remaining Amount: {remaining}")
            return remaining

        def save_payment_details():
            total = int(total_entry.get())
            paid = int(paid_entry.get())
            remaining = calculate_remaining()

            conn = sqlite3.connect('patients.db')
            cursor = conn.cursor()

        # Insert payment details with the current date
            cursor.execute("INSERT INTO payments (patient_id, total, paid, remaining, date) VALUES (?, ?, ?, ?, ?)", 
                       (patient_id, total, paid, remaining, datetime.now().strftime('%Y-%m-%d')))
        
            conn.commit()
            conn.close()

        # Refresh payment history after saving
            display_payment_history()

        def display_payment_history():
            for widget in history_frame.winfo_children():
                widget.destroy()  # Clear previous history

            conn = sqlite3.connect('patients.db')
            cursor = conn.cursor()
            cursor.execute("SELECT date, total, paid, remaining FROM payments WHERE patient_id=? ORDER BY date DESC", (patient_id,))
            payment_history = cursor.fetchall()
            conn.close()

        # Display payment history in the right frame
            for record in payment_history:
                date_label = ctk.CTkLabel(history_frame, text_color="#db3333" ,  text=f"Date: {record[0]}", font=("Dubai", 18, "bold"))
                details_label = ctk.CTkLabel(history_frame,  text=f"Total: {record[1]} | Paid: {record[2]} | Remaining: {record[3]}", font=("Dubai", 16))
                date_label.pack(anchor="w", padx=5, pady=2)
                details_label.pack(anchor="w", padx=5, pady=2)

    # Frames setup
        left_frame = ctk.CTkFrame(tab6 , fg_color="#aff1fc")
        left_frame.pack(side="left", padx=20, pady=20, fill="y")

        right_frame = ctk.CTkFrame(tab6 ,  fg_color="#aff1fc")
        right_frame.pack(side="right", padx=20, pady=20, fill="both", expand=True)

    # Add a title to the Payments tab
        ctk.CTkLabel(left_frame, text="Manage Payments", text_color="#db3333", font=("Dubai", 22, "bold")).pack(pady=20)

    # Create Labels and Entries for Payment Details
        ctk.CTkLabel(left_frame, text="Total Payment:", font=("Dubai", 16)).pack(pady=5 , padx = 20)
        total_entry = ctk.CTkEntry(left_frame, width=200)
        total_entry.pack(pady=5, padx = 10)
        total_entry.insert(0, "0")

        ctk.CTkLabel(left_frame, text="Paid Amount:", font=("Dubai", 16)).pack(pady=5)
        paid_entry = ctk.CTkEntry(left_frame, width=200)
        paid_entry.pack(pady=5, padx = 10)
        paid_entry.insert(0, "0")

    # Display the calculated remaining amount as a label
        remaining_label = ctk.CTkLabel(left_frame, text="Remaining Amount: 0", font=("Dubai", 16))
        remaining_label.pack(pady=5, padx = 10)

    # Update remaining amount whenever the total or paid amount changes
        total_entry.bind("<KeyRelease>", lambda event: calculate_remaining())
        # total_entry.bind("<Returned>", lambda event: calculate_remaining())
        paid_entry.bind("<KeyRelease>", lambda event: calculate_remaining())

    # Save Button
        save_button = ctk.CTkButton(left_frame, fg_color="#00c2ff" ,  text="Save Payment Details", command=save_payment_details)
        save_button.pack(pady=20)

    # History Frame
        history_frame = ctk.CTkFrame(right_frame ,  fg_color="#aff1fc")
        history_frame.pack(padx=10, pady=10, fill="both", expand=True)

    # Calculate and display the initial remaining amount
        calculate_remaining()

    # Display payment history on the right frame
        display_payment_history()

    tab6 = notebook.add("Payments")

# Call the function to add the payment management interface to tab6
    manage_payments()



# Display existing payment details
    payment_display = ctk.CTkLabel(tab6, text="")
    payment_display.pack(pady=20)

    # TAB 6 - Delete Patient
    tab7 = notebook.add("Delete Patient")
    ctk.CTkLabel(tab7, text_color="#db3333", font=("Dubai", 24, "bold"), 
                 text=f"Delete Patient '{patient[1]}' permanently").pack(pady=60)
    delete_button = ctk.CTkButton(tab7, text="Delete Patient", fg_color="#FF1C60", font=("Dubai", 15, "bold"), 
                                  command=lambda: delete_patient(patient_id))
    delete_button.pack()


    def refresh_patient_info():
        global patient  # Use the global patient variable or pass the patient ID to fetch updated info
        updated_patient = fetch_patient(patient[0])
        if updated_patient:
            display_patient_info(left_frame, updated_patient, refresh_patient_info)


    def edit_patient_info(patient, refresh_callback):
    # Create a new window for editing patient info
        edit_window = ctk.CTkToplevel()
        edit_window.title("Edit Patient Info")
        edit_window.geometry("400x800")  # Adjusted height to accommodate additional fields
        edit_window.attributes("-topmost", True)
    
    # Function to save the changes made in the entry fields
        def save_changes():
            new_name = name_entry.get()
            new_age = age_entry.get()
            new_treatment = treatment_entry.get()
            new_phone = phone_entry.get()
            new_gender = gender_var.get()

        # Validate the inputs
            if not new_name or not new_age or not new_treatment or not new_phone or not new_gender:
                ctk.CTkLabel(edit_window, text="All fields are required!", text_color="red").pack(pady=5)
                return

            try:
                new_age = int(new_age)  # Ensure age is an integer
            except ValueError:
                ctk.CTkLabel(edit_window, text="Age must be a number!", text_color="red").pack(pady=5)
                return

        # Update patient information in the database
            conn = sqlite3.connect('patients.db')
            cursor = conn.cursor()
            cursor.execute("UPDATE patients SET name=?, age=?, disease=?, phone=?, gender=? WHERE id=?", 
                       (new_name, new_age, new_treatment, new_phone, new_gender, patient[0]))
            conn.commit()
            conn.close()

            edit_window.destroy()  # Close the edit window
            refresh_callback()  # Refresh the patient info view

    # Create entry fields for patient information
        ctk.CTkLabel(edit_window, text="Name:").pack(pady=5)
        name_entry = ctk.CTkEntry(edit_window)
        name_entry.insert(0, patient[1])
        name_entry.pack(pady=5)

        ctk.CTkLabel(edit_window, text="Age:").pack(pady=5)
        age_entry = ctk.CTkEntry(edit_window)
        age_entry.insert(0, patient[2])
        age_entry.pack(pady=5)

        ctk.CTkLabel(edit_window, text="Treatment Protocol:").pack(pady=5)
        treatment_entry = ctk.CTkEntry(edit_window)
        treatment_entry.insert(0, patient[3])
        treatment_entry.pack(pady=5)

        ctk.CTkLabel(edit_window, text="Phone Number:").pack(pady=5)
        phone_entry = ctk.CTkEntry(edit_window)
        phone_entry.insert(0, patient[4])
        phone_entry.pack(pady=5)

    # Gender selection
        ctk.CTkLabel(edit_window, text="Gender:").pack(pady=5)
        gender_var = ctk.StringVar(value=patient[5])  # Assuming patient[5] contains the current gender
        gender_menu = ctk.CTkOptionMenu(edit_window, fg_color="#db3333" , button_color="#db3333" ,  variable=gender_var, values=["Male", "Female", "Other"])
        gender_menu.pack(pady=5)




    # Save button to save the selected status
        # save_button = ctk.CTkButton(left_frame, text="Save", command=lambda: save_status(patient_id, status_var.get()))
        # save_button.pack(pady=10)

    # Save button
        ctk.CTkButton(edit_window, text="Save Changes", command=save_changes).pack(pady=10)






    def display_patient_info(left_frame, patient, refresh_callback):
        if patient:
        # Function to create a row of label and value
            def add_row(label_text, value_text, label_color="#db3333", value_color="gray"):
                row_frame = ctk.CTkFrame(left_frame, fg_color="white")
                row_frame.pack(fill='x', pady=5, padx=30)
        
                ctk.CTkLabel(row_frame, text=label_text, text_color=label_color, font=("Dubai", 18, "bold")).pack(side='left', padx=(0, 10))
                ctk.CTkLabel(row_frame, text=value_text, text_color=value_color, font=("Dubai", 18, "bold")).pack(side='left')

        # Clear existing widgets
            for widget in left_frame.winfo_children():
                widget.destroy()


            logo_image = Image.open("Images/logo.png")
            logo_image = logo_image.resize((175, 130), Image.LANCZOS)  # Resize the image (width, height)
            logo_photo = ImageTk.PhotoImage(logo_image)
            logo_label = ctk.CTkLabel(left_frame, text="", image=logo_photo)
            logo_label.pack(padx=30, pady=(40 , 30))




            ctk.CTkButton(left_frame, text="Back to Patient List", fg_color="#FF1C60", font=("Dubai", 15, "bold"), command=show_patients).pack(pady=(0 , 30), padx=30)

        # Add patient details to the left frame
            add_row("Patient ID: ", f"{patient[0]:03}")
            add_row("Name: ", f"{patient[1]}")
            add_row("Age: ", f"{patient[2]}")
            add_row("Treatment Protocol: ", f"{patient[3]}")
            add_row("Phone Number: ", f"{patient[4]}")
            add_row("Gender: ", f"{patient[5]}")




# Create the frame and assign it to the variable
            statusFrame = ctk.CTkFrame(left_frame , fg_color="white")
# Use pack()   separately
            statusFrame.pack(pady = 30)

# Create the label and assign it to the variable
            StatusLabel = ctk.CTkLabel(statusFrame, text="Status", font=("dubai", 18, "bold"), text_color="#db3333")
# Use pack() separately
            StatusLabel.pack()
            # global status_var


            status_var = ctk.StringVar(value=patient[6] if patient[6] else "Pending")  # Default value is "Male"
            status_dropdown = ctk.CTkOptionMenu( statusFrame, fg_color="#db3333", button_color="#db3333",  variable=status_var, values=["Pending", "Completed"])
            status_dropdown.pack(padx=10, pady=10 , side = "left")
            btnSave = ctk.CTkButton(statusFrame ,  command=lambda: save_status(patient_id, status_var.get()), text="Save" , width = 25 , font=("dubai" , 13 , "bold")).pack(side = "left")


        # Add a button to go back to the patient list
        
        # Add a button to edit patient info
            ctk.CTkButton(left_frame, text="Edit Info", fg_color="#00c2ff", font=("Dubai", 15, "bold"), command=lambda: edit_patient_info(patient, lambda: display_patient_info(left_frame, fetch_patient(patient[0]), show_patients))).pack(padx=30)




    display_patient_info(left_frame, patient, refresh_patient_info)


    # global back_button


    # Add a "Download All" button

    # ctk.CTkButton(left_frame, text="Back to Patient List", fg_color="#FF1C60", font=("Dubai", 15, "bold"), command=lambda:show_patients).pack(pady=30, padx=30)




def search_patients():
    search_term = search_entry.get().lower()
    for widget in mainFramePat.winfo_children():
        widget.destroy()
    # Adding logos or additional elements
    # Load and display logo at the top center
    # mainFramePat.columnconfigure((0, 1, 2, 3, 4, 5, 6 , 7) , weight=1)

    logo_image = Image.open("Images/FullLogo.png")
    logo_image = logo_image.resize((880, 160), Image.LANCZOS)  # Resize the image (width, height)
    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = ctk.CTkLabel(mainFramePat, text="", image=logo_photo)
    logo_label.grid(row=0, column=0, columnspan=9, pady=10, sticky="n")

    
    # Add a separator line
    line_frame = ctk.CTkFrame(mainFramePat, height=3, fg_color="#db3333")
    line_frame.grid(row=1, column=0, columnspan=9, padx=80, sticky="new")


    # Add widgets to the second frame
    ctk.CTkLabel(mainFramePat, text="Patients List", font=("Dubai", 26, "bold"), text_color="#DB3333").grid(row=1, column=0, columnspan=9, padx=10, pady=20)
    ctk.CTkButton(mainFramePat, text="Back", fg_color="#FF1C60", font=("Dubai", 14, "bold"), command=create_dental_care_window).grid(row=1, column=0, pady=10)

    headers = ["Status" , "Patient-ID", "Name", "Age", "Treatment Protocol", "Phone Number","Gender" ,  "Manage", "Prescribe"]
    for col, header in enumerate(headers):
        ctk.CTkLabel(mainFramePat, text=header, font=("Dubai", 16, "bold")).grid(row=2, column=col, padx= 10 , pady = 30)

    filtered_patients = [
        patient for patient in patients 
        if search_term in patient[1].lower() 
        or search_term == f"{patient[0]:03}"  # Ensuring patient ID is formatted as a 3-digit number
        or search_term in patient[4]
    ]
    for row, patient in enumerate(filtered_patients, start=3):
        patient_id = f"{patient[0]:03}"

        box_color = "#db3333" if patient[6] == "Pending" else "#3333db"

    # Create the colored box (label) for the status
        status_box = ctk.CTkLabel(mainFramePat, text="", width=20, height=20, fg_color=box_color)
        status_box.grid(row=row, column=0, padx=5, pady=20)

        ctk.CTkLabel(mainFramePat, text=patient_id, text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=1, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[1], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=2, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[2], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=3, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[3], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=4, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[4], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=5, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[5], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=6, padx=10, pady=20)
    
        ctk.CTkButton(mainFramePat, text="Manage", font=("Dubai", 14, "bold"), command=lambda p=patient[0]: HistoryPage(p)).grid(row=row, column=7, pady=20, padx=10)
        ctk.CTkButton(mainFramePat, text="Prescribe", fg_color="#db3333", font=("Dubai", 14, "bold"), command=lambda p=patient[0]: PrescriptionPage(p)).grid(row=row, column=8, pady=20, padx=10)

    ctk.CTkButton(mainFramePat, text="Refresh", font=("Dubai", 14, "bold"),  command=show_patients).grid(row=1, column=8, pady=20)



def search_patients_event(event):
    search_patients()





def show_patients():
    # Clear existing widgets
    for widget in root.winfo_children():
        widget.destroy()

    # Create main frame

    global mainFramePat
    mainFramePat = ctk.CTkScrollableFrame(root , fg_color="white")
    mainFramePat.pack(fill=BOTH, expand=1)

    

    mainFramePat.columnconfigure((0, 1, 2, 3, 4, 5, 6 , 7, 8) , weight=1)

    # Adding logos or additional elements
    # Load and display logo at the top center
    logo_image = Image.open("Images/FullLogo.png")
    logo_image = logo_image.resize((880, 160), Image.LANCZOS)  # Resize the image (width, height)
    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = ctk.CTkLabel(mainFramePat, text="", image=logo_photo)
    logo_label.grid(row=0, column=0, columnspan=9, pady=10, sticky="n")

    
    # Add a separator line
    line_frame = ctk.CTkFrame(mainFramePat, height=3, fg_color="#db3333")
    line_frame.grid(row=1, column=0, columnspan=9,  padx=80, sticky="new")
# Load the background image

    

#     # Add widgets to the second frame
    ctk.CTkLabel(mainFramePat, text="Patients List", font=("Dubai", 26, "bold"), text_color="#DB3333").grid(row=1, column=0, columnspan=9, padx=10, pady=20)
    ctk.CTkButton(mainFramePat, text="Back", fg_color="#FF1C60", font=("Dubai", 14, "bold"), command=create_dental_care_window).grid(row=1, column=0, pady=10)

    headers = ["Status", "Patient-ID", "Name", "Age", "Treatment Protocol", "Phone Number", "Gender" ,  "Manage", "Prescribe"]
    for col, header in enumerate(headers):
        ctk.CTkLabel(mainFramePat, text=header, font=("Dubai", 16, "bold")).grid(row=2, column=col, padx= 10 , pady = 30)

    global patients
    patients = fetch_patients()
    for row, patient in enumerate(patients, start=3):
        patient_id = f"{patient[0]:03}"

        box_color = "#db3333" if patient[6] == "Pending" else "#0e8bf1"

    # Create the colored box (label) for the status
        status_box = ctk.CTkLabel(mainFramePat, text="", width=20, height=20, fg_color=box_color)
        status_box.grid(row=row, column=0, padx=5, pady=20)

        ctk.CTkLabel(mainFramePat, text=patient_id, text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=1, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[1], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=2, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[2], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=3, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[3], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=4, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[4], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=5, padx=10, pady=20)
        ctk.CTkLabel(mainFramePat, text=patient[5], text_color="#6F6F6F", font=("Dubai", 15, "bold")).grid(row=row, column=6, padx=10, pady=20)
    
        ctk.CTkButton(mainFramePat, text="Manage", font=("Dubai", 14, "bold"), command=lambda p=patient[0]: HistoryPage(p)).grid(row=row, column=7, pady=20, padx=10)
        ctk.CTkButton(mainFramePat, text="Prescribe", fg_color="#db3333", font=("Dubai", 14, "bold"), command=lambda p=patient[0]: PrescriptionPage(p)).grid(row=row, column=8, pady=20, padx=10)

#     ctk.CTkButton(second_frame, text="Refresh", font=("Dubai", 14), command=show_patients).grid(row=len(patients) + 2, column=6, pady=20)


    # Add search bar
    search_bar_frame = ctk.CTkFrame(mainFramePat , fg_color="white")
    search_bar_frame.grid(row = 1 , column = 7 , columnspan=2)
    global search_entry
    search_entry = ctk.CTkEntry(search_bar_frame, placeholder_text="Search by Name, ID, or Phone...", width=250 ,font=("Dubai", 14))
    search_entry.pack(side=LEFT, padx=10)
    search_entry.bind("<Return>", search_patients_event)  
    ctk.CTkButton(search_bar_frame, text="Search", fg_color="#FF1C60", font=("Dubai", 14, "bold"), command=search_patients).pack(pady = 10)
    






def fetch_patient(patient_id):
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name, age, disease , phone , gender , status FROM patients WHERE id=?", (patient_id,))
    row = cursor.fetchone()
    conn.close()
    return row

def fetch_prescriptions(patient_id):
    conn = sqlite3.connect('patients.db')
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM prescriptions WHERE patient_id=?", (patient_id,))
    rows = cursor.fetchall()
    conn.close()
    return rows

# Function to fetch the saved medical history text


def backside():
    # Clear existing widgets
    for widget in root.winfo_children():
        widget.destroy()

    # Main Frame
    mainFrameBackside = ctk.CTkFrame(root, fg_color="white")
    mainFrameBackside.pack(fill=tk.BOTH, expand=1)

    # Left Side Frame
    left_side2 = ctk.CTkScrollableFrame(
        mainFrameBackside,
        fg_color="white",
        scrollbar_fg_color="white",
        scrollbar_button_color="#b0f3ff",
        scrollbar_button_hover_color="green",
        width=280
    )
    left_side2.pack(side="left", fill=tk.Y)

    # Logo Image
    logo_image = Image.open("Images/logo.png")
    logo_image = logo_image.resize((175, 130), Image.LANCZOS)
    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = ctk.CTkLabel(left_side2, text="", image=logo_photo)
    logo_label.pack(padx=30, pady=40)

    # Entry Fields
    ctk.CTkLabel(left_side2, text="Treatment:", fg_color="white").pack(padx=10, pady=5)
    global treatment_entry
    treatment_entry = ctk.CTkEntry(left_side2, width=200)
    treatment_entry.pack(padx=10, pady=5)

    ctk.CTkLabel(left_side2, text="Date (YYYY-MM-DD):", fg_color="white").pack(padx=10, pady=5)
    global date_entry
    date_entry = ctk.CTkEntry(left_side2, width=200)
    date_entry.pack(padx=10, pady=5)

    ctk.CTkLabel(left_side2, text="Time (HH:MM) AM/PM:", fg_color="white").pack(padx=10, pady=5)
    global time_entry
    time_entry = ctk.CTkEntry(left_side2, width=200)
    time_entry.pack(padx=10, pady=5)


    # Middle Frame
    middle_frame2 = ctk.CTkFrame(mainFrameBackside, fg_color="#b0f3ff")
    middle_frame2.pack(side="left", fill="both", expand=1)

    ctk.CTkButton(middle_frame2 , text="Print" , fg_color="#db3333" , command=generatebackpdf).pack(side = "bottom" , pady  = 10)
    previewpageframe2 = ctk.CTkFrame(middle_frame2, fg_color="white", height=1100)
    previewpageframe2.pack(fill=tk.Y, pady=30, expand=1)


    # Logo Image in Preview Page Frame
    logo_image2 = Image.open("Images/backside.jpeg")
    logo_image2 = logo_image2.resize((500,550), Image.LANCZOS)
    logo_photo2 = ImageTk.PhotoImage(logo_image2)
    logo_label2 = ctk.CTkLabel(previewpageframe2, text="", image=logo_photo2)
    logo_label2.pack(padx=10, pady=10)

    # Static Labels
    ctk.CTkLabel(previewpageframe2, text="Next Appointment", text_color="#db3333", font=("dubai", 18, "bold")).pack()

    anotherFrame = ctk.CTkFrame(previewpageframe2, fg_color="white")
    anotherFrame.pack(side="left", fill=tk.BOTH, expand=1, pady= 10)
    anotherFrame.columnconfigure((0, 1), weight=1)

    # Labels for dynamic content
    global a, b, c
    a = ctk.CTkLabel(anotherFrame, text="", text_color="black", font=("dubai", 15))
    a.grid(column=0, row=1)

    b = ctk.CTkLabel(anotherFrame, text="", text_color="black", font=("dubai", 15))
    b.grid(column=1, row=1)

    c = ctk.CTkLabel(anotherFrame, text="", text_color="black", font=("dubai", 15))
    c.grid(column=1, row=3)

    ctk.CTkLabel(anotherFrame, text="Treatment:", text_color="#db3333", font=("dubai", 16, "bold")).grid(column=0, row=0)
    ctk.CTkLabel(anotherFrame, text="Date:", text_color="#db3333", font=("dubai", 16, "bold")).grid(column=1, row=0)
    ctk.CTkLabel(anotherFrame, text="Time:", text_color="#db3333", font=("dubai", 16, "bold")).grid(column=1, row=2)

    # ctk.CTkLabel(anotherFrame, text="See you Next Time", text_color="#db3333", font=("dubai", 14, "bold")).grid(column=0, row=4 , pady=(0 , 10),  columnspan = 2)

    # Function to display information
    def display_info():
        treatment = treatment_entry.get()
        date = date_entry.get()
        time = time_entry.get()

        a.configure(text=f"{treatment}")
        b.configure(text=f"{date}")
        c.configure(text=f"{time}")

    # Add Display Button
    display_button = ctk.CTkButton(left_side2, text="Display Info", command=display_info)
    display_button.pack(padx=10, pady=20)

    ctk.CTkButton(left_side2, text="Back to Patient List",font=("dubai" , 13 , "bold") ,fg_color="#db3333", hover_color="red", width=150,command=show_patients).pack()




    


    # # Static Labels
    # static_labels = [
    #     ctk.CTkLabel(previewpageframe2, text="Next Appointment", text_color="#db3333", font=("dubai", 18, "bold")),
    #     # ctk.CTkLabel(previewpageframe2, text="See You Next Time!", text_color="#db3333", font=("dubai", 18, "bold"))
    # ]
    # for label in static_labels:
    #     label.pack()







def PrescriptionPage(patient_id):
    for widget in root.winfo_children():
        widget.destroy()

    mainFramePrescribe = ctk.CTkFrame(root, fg_color="white")
    mainFramePrescribe.pack(fill=BOTH, expand=1)

    left_side = ctk.CTkScrollableFrame(mainFramePrescribe, fg_color="white" , scrollbar_fg_color="white" , scrollbar_button_color="#b0f3ff" , scrollbar_button_hover_color="green" ,  width = 280 )
    left_side.pack(side="left", fill=Y)

    logo_image = Image.open("Images/logo.png")
    logo_image = logo_image.resize((175, 130), Image.LANCZOS)  # Resize the image (width, height)
    logo_photo = ImageTk.PhotoImage(logo_image)
    logo_label = ctk.CTkLabel(left_side, text="", image=logo_photo)
    logo_label.pack(padx=30, pady=40)

    ctk.CTkButton(left_side, text="Back to Patient List",font=("dubai" , 13 , "bold") ,fg_color="#db3333", hover_color="red", width=150,command=show_patients).pack()


    ctk.CTkLabel(left_side, text="Medicines:", text_color="#db3333", font=("dubai", 22, "bold")).pack(padx=50, pady=20)

    medicines = fetch_medicines()

    global selected_medicines


    for idx, (med_id, med_name) in enumerate(medicines):
        frame = ctk.CTkFrame(left_side , fg_color="white")
        frame.pack(pady=6,  fill=BOTH)
        ctk.CTkButton(frame, text=med_name, text_color="white", bg_color="white", font=("Dubai", 15, "bold"), command=lambda m=med_name: add_medicine(m)).pack(side="left" , padx = (10, 0))
        ctk.CTkButton(frame, text="X" , width=40 ,  text_color="white", bg_color="white", font=("Dubai", 15, "bold"), fg_color="#db3333", command=lambda m_id=med_id: delete_medicine(m_id)).pack(padx=(10 , 4) , anchor = "e")





    middle_frame = ctk.CTkFrame(mainFramePrescribe, fg_color="#b0f3ff")
    middle_frame.pack(side="left", fill="both", expand=1)

    ctk.CTkButton(middle_frame , text="Print" , fg_color="#db3333" , command=lambda:save_prescription_and_generate_pdf(patient_id)).pack(side="bottom" , pady=20)



    previewpageframe = ctk.CTkFrame(middle_frame , fg_color="white" , height=1100)
    previewpageframe.pack(fill = Y , pady = 30,  expand = 1)




    logo_image2 = Image.open("Images/FullLogo.png")
    logo_image2 = logo_image2.resize((600, 90), Image.LANCZOS)  # Resize the image (width, height)
    logo_photo2 = ImageTk.PhotoImage(logo_image2)
    logo_label2 = ctk.CTkLabel(previewpageframe, text="", image=logo_photo2)
    logo_label2.pack(padx = 10 , pady = 10)
    
    # Add a separator line
    line_frame = ctk.CTkFrame(previewpageframe, height=3, width=330,fg_color="#db3333")
    line_frame.pack(padx = 10 )

    logo_image3 = Image.open("Images/doctors.png")
    logo_image3 = logo_image3.resize((600, 130), Image.LANCZOS)  # Resize the image (width, height)
    logo_photo3 = ImageTk.PhotoImage(logo_image3)
    logo_label3 = ctk.CTkLabel(previewpageframe, text="", image=logo_photo3)
    logo_label3.pack(padx = 10 , pady = 10)




    def open_medical_history(patient_id):
        def add_checked_values():
            selected_items = []
            for checkbox, var in checkboxes:
                if var.get() == 1:
                    selected_items.append(checkbox.cget('text'))

            # Update the global selected_medical_history
            global selected_medical_history

            # Append new selections to the previous selections
            if selected_medical_history:
                selected_medical_history += '\n' + '\n'.join(selected_items)
            else:
                selected_medical_history = '\n'.join(selected_items)

            # Display selected items under the "Medical History" label
            medical_history_label.configure(text=selected_medical_history)

            # Save the updated medical history to the database
            save_medical_checked_history(patient_id)

            # Append the new history to the existing one in tab2
            previous_history = history_label.cget("text")
            new_history = f"{selected_medical_history}\n{previous_history}"
            history_label.configure(text=new_history)

        # Create the medical history window
        medical_window = ctk.CTkToplevel(root)
        medical_window.title("Medical History")
        medical_window.geometry("300x400")
        medical_window.attributes("-topmost", True)

        # Fetch all available general medical history
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("SELECT history FROM general_medical_history")
        histories = cursor.fetchall()

        # Fetch saved medical history for the patient
        cursor.execute("SELECT SavedMedicalHistory FROM SaveMedicalHistoryChecked WHERE patient_id=? ORDER BY date DESC", (patient_id,))
        saved_history = cursor.fetchone()
        conn.close()

        # Convert saved history to a list for easier comparison
        if saved_history:
            saved_items = saved_history[0].split('\n')
        else:
            saved_items = []

        checkboxes = []
        for history in histories:
            history_list = history[0].split(",")
            for item in history_list:
                item = item.strip()
                var = ctk.IntVar(value=1 if item in saved_items else 0)
                checkbox = ctk.CTkCheckBox(medical_window, text=item, variable=var)
                checkbox.pack(padx=10, pady=5)
                checkboxes.append((checkbox, var))

        # Add a button to add checked values
        add_button = ctk.CTkButton(medical_window, text="Add to Medical History", command=add_checked_values)
        add_button.pack(pady=10)




    def open_dental_history(patient_id):
        def add_checked_values():
            selected_items = []
            for checkbox, var in checkboxes:
                if var.get() == 1:
                    selected_items.append(checkbox.cget('text'))

            # Display selected items under the "Dental History" label
            global selected_dental_history
            selected_dental_history = '\n'.join(selected_items)
            dental_history_label.configure(text=selected_dental_history)
            root.update_idletasks()

            # Save the selected items to the database
            save_dental_checked_history(patient_id)

            # Append the new history to the existing one in tab3
            previous_history = history_label2.cget("text")
            new_history = f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}:\n{selected_dental_history}\n\n{previous_history}"
            history_label2.configure(text=new_history)

        # Fetch previously checked values
        previous_dental_history = fetch_dental_checked_history(patient_id)

        dental_window = ctk.CTkToplevel(root)
        dental_window.title("Dental History")
        dental_window.geometry("300x400")
        dental_window.attributes("-topmost", True)

        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("SELECT history FROM general_dental_history")
        histories = cursor.fetchall()
        conn.close()

        checkboxes = []
        for history in histories:
            history_list = history[0].split(",")
            for item in history_list:
                var = ctk.IntVar()
                checkbox = ctk.CTkCheckBox(dental_window, text=item.strip(), variable=var)
                checkbox.pack(padx=10, pady=5)

                # Check if this item was previously selected
                if item.strip() in previous_dental_history:
                    var.set(1)

                checkboxes.append((checkbox, var))

        # Add a button to add checked values
        add_button = ctk.CTkButton(dental_window, text="Add to Dental History", command=add_checked_values)
        add_button.pack(pady=10)



    def open_treatment_planning():
        def add_checked_values():
            selected_items = []
            for checkbox, var in checkboxes:
                if var.get() == 1:
                    selected_items.append(checkbox.cget('text'))

        # Display selected items under the "Dental History" label in a vertical list
            global selected_treatment_history
            selected_treatment_history = '\n'.join(selected_items)
            treatment_history_label.configure(text=selected_treatment_history)

        treatment_window = ctk.CTkToplevel(root)
        treatment_window.title("Treatment Planning History")
        treatment_window.geometry("300x400")
        treatment_window.attributes("-topmost", True)

        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("SELECT history FROM treatment_planning")
        histories = cursor.fetchall()
        conn.close()

        checkboxes = []
        for history in histories:
            history_list = history[0].split(",")
            for item in history_list:
                var = ctk.IntVar()
                checkbox = ctk.CTkCheckBox(treatment_window, text=item.strip(), variable=var)
                checkbox.pack(padx=10, pady=5)
                checkboxes.append((checkbox, var))

    # Add a button to add checked values
        add_button = ctk.CTkButton(treatment_window, text="Add to Treatment Planning", command=add_checked_values)
        add_button.pack(pady=10) 




    detailsframe = ctk.CTkFrame(previewpageframe , fg_color="white")
    detailsframe.pack(side = "top" , fill = X)
    detailsframe.columnconfigure((0, 1, 2, 3) , weight=1)
    patient = fetch_patient(patient_id)

    current_date = date.today().strftime("%B %d, %Y")
    ctk.CTkLabel(detailsframe , font=("dubai" , 12 , "bold")  , text = f"Date: {current_date}").grid(row = 0 , column = 0)
    ctk.CTkLabel(detailsframe , font=("dubai" , 12 , "bold")  , text = f"Pt. Name: {patient[1]}").grid(row = 0 , column = 1 )
    ctk.CTkLabel(detailsframe , font=("dubai" , 12 , "bold")  , text = f"Age: {patient[2]}").grid(row = 0 , column = 2)
    ctk.CTkLabel(detailsframe , font=("dubai" , 12 , "bold")  , text = f"Ref: {patient[0]:03}").grid(row = 0 , column = 3)


#here
    secFrame = ctk.CTkFrame(previewpageframe , fg_color= "white")
    secFrame.pack(side = "left" , fill = BOTH , expand = 1)

# Create the main frame
    LeftsecFrame = ctk.CTkFrame(secFrame, fg_color="white")
    LeftsecFrame.pack(side="left", fill= Y)

# Create the border frame on the right side
    border_frame = ctk.CTkFrame(
    LeftsecFrame, 
    width=3,                 # Set the width of the border
    height = 100,  # Match the height of the main frame
    fg_color="black"         # Set the color of the border
)
    border_frame.pack(side="right", fill=Y)


    ctk.CTkLabel(LeftsecFrame , font=("dubai" , 14 , "bold")  , text = f"Medical History").pack(padx = 20 , pady = 10)
    # Assuming detailsframe is already created as per your code
    medical_history_label = ctk.CTkLabel(LeftsecFrame, text_color="gray" ,font=("dubai", 12, "bold"), text="")
    medical_history_label.pack()

    ctk.CTkLabel(LeftsecFrame , font=("dubai" , 14 , "bold")  , text = f"Dental History").pack(padx = 20 , pady = (40 , 0))
    dental_history_label = ctk.CTkLabel(LeftsecFrame, font=("dubai", 12,"bold"), text_color="gray" , text="")
    dental_history_label.pack()




    RightsecFrame = ctk.CTkFrame(secFrame , fg_color= "white")
    RightsecFrame.pack(side = "left" , fill = BOTH , expand = 1)
    # ctk.CTkButton(middle_frame , text="Print").pack(side="bottom" , anchor = "w")
    # RightsecFrame.columnconfigure((0,1) , weight=1)



# Initialize the label in the detailsframe
    medicine_label = ctk.CTkLabel(RightsecFrame, text="" , text_color="gray" , font=("dubai" , 12 , "bold"))
    medicine_label.pack(side = "left" , anchor ="n" , pady =10 , padx = 10)



    ctk.CTkLabel(RightsecFrame , font=("dubai" , 14 , "bold")  , text = f"Treatment Planning").pack(padx = (30, 10) , pady = (170 , 0))
    treatment_history_label = ctk.CTkLabel(RightsecFrame, font=("dubai", 12,"bold"), text_color="gray" , text="")
    treatment_history_label.pack()

    # ctk.CTkButton(middle_frame , text="print" , command=save_prescription_and_generate_pdf).pack()

    right_frame = ctk.CTkFrame(mainFramePrescribe, fg_color="white")
    right_frame.pack(side="right", fill="y")

    AddNewMedicineFrame = ctk.CTkFrame(right_frame, fg_color="#b0f3ff")
    AddNewMedicineFrame.pack(padx=50, pady=20)

    def add_new_medicine():
        new_medicine = new_medicine_entry.get()
        if new_medicine and new_medicine not in [med[1] for med in medicines]:  # Check if medicine is not already in the list
            conn = sqlite3.connect('patients.db')
            cursor = conn.cursor()
            cursor.execute("INSERT INTO medicines (name) VALUES (?)", (new_medicine,))
            conn.commit()
            conn.close()
            PrescriptionPage(patient_id)

    # New medicine entry and button on the right
    ctk.CTkLabel(AddNewMedicineFrame,text_color="#db3333" , text="Add New Medicine", font=("dubai", 16 , "bold")).pack(padx=10, pady=20)
    new_medicine_entry = ctk.CTkEntry(AddNewMedicineFrame, placeholder_text="Enter medicine name")
    new_medicine_entry.pack(pady=20, padx=20)
    new_medicine_entry.bind("<Return>", lambda event: add_new_medicine())
    ctk.CTkButton(AddNewMedicineFrame, text="Add", text_color="white", font=("Dubai", 15, "bold"), command=add_new_medicine).pack(pady=(0, 10), padx=20)

  

    # Medical History
    medical_history_frame = ctk.CTkFrame(right_frame, fg_color="white")
    medical_history_frame.pack(padx=50, pady=20 )
    medical_HistoryBTN = ctk.CTkButton(medical_history_frame,text_color="white",
        font=("Dubai", 13, "bold"),
      command=lambda:open_medical_history(patient_id), text="Select Medical History").pack(pady=10, padx=20)
    dental_HistoryBTN = ctk.CTkButton(medical_history_frame,text_color="white",
        font=("Dubai", 13, "bold"),
         command=lambda:open_dental_history(patient_id), text="Select Dental History").pack(pady=10, padx=20)
    TreatmentPlanningBTN = ctk.CTkButton(medical_history_frame,text_color="white",
        font=("Dubai", 13, "bold"),
        command=open_treatment_planning, text="Select Treatment Planning").pack(pady=10, padx=20)


    backBTN = ctk.CTkButton(right_frame , fg_color="#db3333" ,  text="Back Side" , font=("dubai" , 15 , "bold") , command=backside).pack()





  
    selected_medicines = {}  # Dictionary to store medicines and their "checked" status

    def open_medicine_window():
        medicines = [med for med in fetch_medicines() if med[1] in selected_medicines]

    # Create a new window
        medicine_window = ctk.CTkToplevel()
        medicine_window.title("Set Medicine Timing")
        medicine_window.attributes("-topmost", True)

        def on_checkbox_toggle():
            selected_text = []
            for id, name in medicines:
                if selected_medicines[name].get():
                    selected_text.append(f"{name}\t\t صبح  شام  ")
                else:
                    selected_text.append(name)
            medicine_label.configure(text="\n".join(selected_text))

    # Create checkboxes for each selected medicine
        for id, name in medicines:
            selected_medicines[name] = ctk.BooleanVar(value=False)  # Initialize with True if pre-selected
            checkbox = ctk.CTkCheckBox(medicine_window, text=name, variable=selected_medicines[name], command=on_checkbox_toggle)
            checkbox.pack(anchor="w", padx=20, pady=5)  # Aligning checkboxes to the left with padding

    # Add an 'OK' button to confirm selection
        confirm_button = ctk.CTkButton(medicine_window, text="OK", text_color="white", fg_color="#00c2ff", command=lambda m_id=med_id: [medicine_window.destroy()])
        confirm_button.pack(pady=10)


# New button to open the popup window
    set_medicine_button = ctk.CTkButton(
        medical_history_frame,
        text="Set Medicine Timing",
        text_color="white",
        font=("Dubai", 13, "bold"),
        command=open_medicine_window  # Link button to open the medicine window
)
    set_medicine_button.pack(pady=10, padx=20)        

    def update_medicine_list():
    # Update the label text with the selected medicines and "checked" status
        selected_text = []
        for name, var in selected_medicines.items():
            if var.get():
                selected_text.append(f"{name} \t\t صبح  شام")
            else:
                selected_text.append(name)
        medicine_label.configure(text="\n".join(selected_text))

    def add_medicine(medicine):
        if medicine not in selected_medicines:
            selected_medicines[medicine] = ctk.BooleanVar(value=False)  # Initialize as checked
            update_medicine_list()  # Update the displayed list

    # def remove_medicine():
    #     selected = medicine_listbox.get(medicine_listbox.curselection())
    #     if selected in selected_medicines:
    #         selected_medicines.remove(selected)
    #         update_medicine_list()

    # def update_medicine_list():
    #     medicine_listbox.delete(0, END)
    #     for medicine in selected_medicines:
    #         medicine_listbox.insert(END, medicine)

    def delete_medicine(medicine_id):
        conn = sqlite3.connect('patients.db')
        cursor = conn.cursor()
        cursor.execute("DELETE FROM medicines WHERE id=?", (medicine_id,))
        conn.commit()
        conn.close()
        PrescriptionPage(patient_id)

    # Medicine listbox and other controls on the right
    # medicine_listbox = Listbox(root, fg="gray", font=("dubai", 22, "bold"), bd=2, relief="solid", highlightbackground="light blue", highlightcolor="blue", highlightthickness=2)
    # medicine_listbox.grid(row=1, column=1, columnspan=1, rowspan=5, pady=5, padx=10, sticky="news")
    # ctk.CTkButton(root, text="Remove Selected", text_color="white", bg_color="white", font=("Dubai", 15, "bold"), fg_color="#db3333", command=remove_medicine).grid(row=6, column=1, pady=5, sticky="nw")
    # ctk.CTkButton(root, text="Print Prescription", text_color="white", bg_color="white", font=("Dubai", 15, "bold"), fg_color="#00c2ff", command=lambda: save_prescription_and_generate_pdf(patient_id)).grid(row=6, column=1, pady=5, sticky="ne")  # Added print button
    # ctk.CTkButton(root, text="Back to Patient List", text_color="white", bg_color="white", font=("Dubai", 15, "bold"), fg_color="#db3333", command=show_patients).grid(row=10, column=2, pady=10)




# Function to show the main window (dental care window)
def show_main_window():
    welcome_window.destroy()

# Create the welcome window
welcome_window = ctk.CTk()
welcome_window.title("Welcome")
# Set background color to match the main window color (change to your desired color)
bg_color = "white"  # Example color, adjust as needed
welcome_window.configure(bg=bg_color)
# Maximize the welcome window
welcome_window.attributes("-fullscreen", True)

# Center the content in the welcome window


# # Display "Welcome" message
welcome_label = ctk.CTkLabel(welcome_window, text="Welcome", font=("Dubai", 38, "bold"), text_color="#db3333")
welcome_label.pack(pady=30, padx=10)
# ctk.CTkLabel(welcome_window, text_color="#db3333",  text="Welcome" , font=("dubai" , 24 ,  "bold"))
# Display company logo (replace 'Images/windowloading.png' with the path to your logo file)
logo_image = Image.open("Images/windowloading.png")
logo_photo = ImageTk.PhotoImage(logo_image)
logo_label = ctk.CTkLabel(welcome_window, text="", image=logo_photo, anchor="center")
logo_label.pack(pady=10, padx=10, expand=True, fill='both')

# # Display "NH Clinic Management System" text
nh_label = ctk.CTkLabel(welcome_window, text="NH Clinic Management System", font=("Dubai", 28), text_color="#db3333")
nh_label.pack(pady=30, padx=10)

# # Set a timer to show the main window after 20 seconds
welcome_window.after(1500, show_main_window)

welcome_window.mainloop()

def play_video(video_path, on_complete):
    video = imageio.get_reader(video_path)
    try:
        fps = video.get_meta_data()['fps']
        skip_frames = 2  # Skip every other frame for 2x speed
        for i, frame in enumerate(video.iter_data()):
            if i % skip_frames == 0:
                # Resize the frame to match the root window's size
                frame_image = Image.fromarray(frame)
                resized_frame = frame_image.resize((root.winfo_width(), root.winfo_height()), Image.LANCZOS)
                frame_image_tk = ImageTk.PhotoImage(resized_frame)
                if video_label.winfo_exists():
                    root.after(0, update_video_label, frame_image_tk)
                    root.after(int(1000 / (fps * 2)))  # Adjust timing to account for 2x speed
        if video_label.winfo_exists():
            root.after(0, on_complete)  # Call the on_complete function when the video ends
    except Exception as e:
        print(f"Error playing video {video_path}: {e}")

def update_video_label(frame_image):
    if video_label.winfo_exists():  # Check if video_label exists before updating
        video_label.configure(image=frame_image)
        video_label.image = frame_image  # Keep a reference to avoid garbage collection

def initialize_welcome_window():
    for widget in root.winfo_children():
        widget.destroy()

    # Create the welcome window
    welcome_window2 = ctk.CTkFrame(root)
    welcome_window2.configure(fg_color="black")  # Use black background during video playback
    welcome_window2.pack(fill="both", expand=True)
    root.attributes('-fullscreen', True)
 

    # Video playback label (fullscreen)
    global video_label
    video_label = ctk.CTkLabel(welcome_window2, text="")
    video_label.pack(fill='both', expand=True)  # Ensure the video label covers the entire window

    global text_widget, logo_label, images, texts, body_texts, current_index, videos
    videos = ["videos/1.mp4", "videos/2.mp4", "videos/3.mp4",
              "videos/4.mp4", "videos/5.mp4", "videos/6.mp4"]
    images = ["Images/image1.jpg", "Images/image2.jpg", "Images/image3.jpg",
              "Images/image4.jpg", "Images/image5.jpg", "Images/image6.jpg"]
    texts = [
        "Welcome to NH Dental Care Clinic!",
        "We provide excellent care.",
        "Meet Our Expert Team",
        "Schedule your appointment today.",
        "Experience the best service.",
        "Thank you for visiting!"
    ]
    body_texts = [
        "Welcome to NH Dental Care Clinic, a trusted name in dental care. Our clinic has been dedicated to providing high-quality dental services to the community. Our journey began with a vision to create a welcoming and comfortable environment where patients of all ages can receive exceptional dental care.",
        "We offer a wide range of dental services to meet your needs.",
        "Dr. Zunaira - Dental Surgeon\tQualifications: B.D.S, R.D.S\n\nDr. Beena Shadab Khan - Maxillofacial Surgeon\tQualifications: B.D.S, R.D.S, MDS, (PGR) OMFS\n\nDr. Abdul Nasir Khan - Dental Surgeon\tQualifications: B.D.S, R.D.S, B.Sc (Eng)\n\nDr. Laraib Shoukat - Dental Surgeon\tQualifications: B.D.S, R.D.S\n\nDr. Hameez Sultan - Dental Surgeon\tQualifications: B.D.S, R.D.S",
        "Book your appointment now and take the first step towards a healthier smile.",
        "Here, we pride ourselves on our state-of-the-art facilities, equipped with the latest dental technology for precise diagnostics and effective treatments. Our comprehensive services range from routine check-ups to advanced cosmetic procedures, tailored to meet all your dental needs. We prioritize personalized care, treating each patient individually to achieve specific dental health goals. Our friendly staff and serene environment ensure comfort and convenience, with flexible scheduling to accommodate your busy lifestyle. Committed to excellence, our dental professionals continuously learn and stay updated with the latest advancements in dentistry.",
        "We appreciate your trust in us and look forward to seeing you!"
    ]

    current_index = 0  # Start index for cycling through lists

    def start_video_loop():
        def play_videos():
            if video_label.winfo_exists():
                for video_path in videos:
                    play_video(video_path, lambda: None)
                if welcome_window2.winfo_exists():
                    root.after(0, initialize_text_and_images)  # After videos, initialize content

        Thread(target=play_videos).start()

    def initialize_text_and_images():
        if welcome_window2.winfo_exists():
            welcome_window2.configure(fg_color="white") 
            
             # Change background to white for content
            video_label.pack_forget()  # Hide the video label

            # Create frames for left (image) and right (logo & text)
            content_frame = ctk.CTkFrame(welcome_window2, fg_color="white")
            content_frame.pack(fill="both", expand=True)

            # Left frame for image
            left_frame = ctk.CTkFrame(content_frame, fg_color="white")
            left_frame.pack(side="left", fill="both", expand=True)

            # Right frame for logo and text
            right_frame = ctk.CTkFrame(content_frame, fg_color="white")
            right_frame.pack(side="right", fill="both", expand=True)

            # Inside right frame: create top section for logo and bottom section for text
            logo_frame = ctk.CTkFrame(right_frame, fg_color="white")
            logo_frame.pack(side="top", fill="x", padx=20, pady=10)

            text_frame = ctk.CTkFrame(right_frame, fg_color="white")
            text_frame.pack(side="top", fill="both", expand=True, padx=20, pady=10)

            # Image on the left side
            global logo_label
            image_label = ctk.CTkLabel(left_frame, text="")
            image_label.pack(pady=10, padx=10, expand=True, fill="both")

            # Logo at the top right
            logo_label = ctk.CTkLabel(logo_frame, text="")
            logo_label.pack(expand=True, fill="both")

            # Text on the bottom right
        # Heading text (red color)
            global heading_label
            heading_label = ctk.CTkLabel(text_frame, text="", text_color="red", font=("dubai", 18, "bold"))
            heading_label.pack(anchor="nw", pady=20, padx = 20)  # Top alignment with padding below

        # Body text (black color)
            global body_label
            body_label = ctk.CTkLabel(text_frame, text="", text_color="black", font=("dubai", 18) ,justify="left", wraplength=600 )
            body_label.pack(anchor="w", pady=20 ,padx = 20)  # Expands to fill available space

            # Populate the text widget with content
            update_content(image_label)  # Start content loop with the image label

    def update_content(image_label):
        global current_index

        try:
            if heading_label.winfo_exists() and body_label.winfo_exists():
            # Update the heading and body text
                heading_label.configure(text=texts[current_index])
                body_label.configure(text=body_texts[current_index])

                IMAGE_WIDTH = 800
                IMAGE_HEIGHT = 800
                new_image_photo = resize_image(images[current_index], IMAGE_WIDTH, IMAGE_HEIGHT)
                if new_image_photo and image_label.winfo_exists():
                    image_label.configure(image=new_image_photo)
                    image_label.image = new_image_photo  # Keep a reference to avoid garbage collection

                # Set logo at the top right (persistent)
                logo_photo = resize_image("Images/fullLogo.png", 1000, 170)  # Adjust logo size
                if logo_photo and logo_label.winfo_exists():
                    logo_label.configure(image=logo_photo)
                    logo_label.image = logo_photo  # Keep reference to avoid garbage collection

                current_index += 1

                if current_index >= len(texts):
                    current_index = 0  # Reset index to loop content again
                    if welcome_window2.winfo_exists():
                        welcome_window2.after(3000, initialize_welcome_window)  # Start video loop again after content
                else:
                    if welcome_window2.winfo_exists():
                        welcome_window2.after(3000, lambda: update_content(image_label))  # Update every 3 seconds
        except Exception as e:
            print(f"Error updating content: {e}")

    root.bind("<Escape>", lambda event: root.quit())
    start_video_loop()  # Start the video loop immediately

def resize_image(image_path, width, height):
    try:
        image = Image.open(image_path)
        resized_image = image.resize((width, height), Image.LANCZOS)
        return ImageTk.PhotoImage(resized_image)  # Convert to PhotoImage
    except Exception as e:
        print(f"Error resizing image: {e}")
        return None


# Main

root = ctk.CTk()
root.title("Patient Management System")
root.geometry("1200x900")
root.iconbitmap('Images/logo.ico')




create_or_alter_table()
# show_patients()
create_dental_care_window()


root.mainloop()
